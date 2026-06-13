#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软件著作权申请表生成脚本
生成标准格式的软件著作权登记申请表文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcPr.append(element)


def set_run_font(run, font_name='宋体', font_size=12, bold=False):
    """设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold


def add_paragraph_with_font(doc, text, font_name='宋体', font_size=12, bold=False, alignment=None):
    """添加段落并设置字体"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    return p


def add_field_row(table, label, value, label_bold=False, value_bold=False):
    """添加字段行"""
    row = table.add_row()
    cells = row.cells

    # 设置标签单元格
    label_para = cells[0].paragraphs[0]
    label_para.clear()
    label_run = label_para.add_run(label)
    set_run_font(label_run, '宋体', 10.5, label_bold)

    # 设置值单元格
    value_para = cells[1].paragraphs[0]
    value_para.clear()
    value_run = value_para.add_run(value)
    set_run_font(value_run, '宋体', 10.5, value_bold)

    return row


def create_application_form():
    """创建软件著作权申请表"""
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    # 标题
    title = add_paragraph_with_font(
        doc,
        '软件著作权登记申请表',
        font_name='黑体',
        font_size=22,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    title.paragraph_format.space_after = Pt(20)

    # 创建表格
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    # 设置列宽
    for cell in table.columns[0].cells:
        cell.width = Inches(2)
    for cell in table.columns[1].cells:
        cell.width = Inches(4.5)

    # 软件全称
    add_field_row(table, '软件全称', '帮信罪辅助裁定系统', label_bold=True, value_bold=True)

    # 软件简称
    add_field_row(table, '软件简称', '帮信裁定系统')

    # 版本号
    add_field_row(table, '版本号', 'V1.0')

    # 软件分类
    add_field_row(table, '软件分类', '应用软件 - 法律工具')

    # 开发完成日期
    add_field_row(table, '开发完成日期', '2026-06-11')

    # 首次发表日期
    add_field_row(table, '首次发表日期', '')

    # 发表状态
    add_field_row(table, '发表状态', '未发表')

    # 权利取得方式
    add_field_row(table, '权利取得方式', '原始取得')

    # 权利范围
    add_field_row(table, '权利范围', '全部权利')

    # 编程语言
    add_field_row(table, '编程语言', 'Python 3.11+ / JavaScript / Vue 3 / SQL')

    # 开发工具
    add_field_row(table, '开发工具', 'VS Code / Trae IDE')

    # 运行环境
    add_field_row(table, '运行环境', '详见设计说明书')

    # 技术特点
    add_field_row(table, '技术特点', '详见设计说明书')

    # 程序总行数
    add_field_row(table, '程序总行数', '（待统计）')

    # 开发者信息
    add_field_row(table, '开发者', '待填')

    add_field_row(table, '联系人', '待填')

    add_field_row(table, '联系电话', '待填')

    add_field_row(table, '电子邮箱', '待填')

    add_field_row(table, '联系地址', '待填')

    # 保存文档
    output_path = '01_软件著作权申请表/软著申请表_帮信罪辅助裁定系统V1.0.docx'
    doc.save(output_path)
    print(f'软件著作权申请表已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_application_form()
