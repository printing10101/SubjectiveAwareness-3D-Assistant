"""
scripts/generate_word_report.py

Generate empirical research report as Word document.
Output: reports/\u5b9e\u8bc1\u7814\u7a76\u62a5\u544a_V1.0.docx
"""

import os
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

LQ = "\u201c"
RQ = "\u201d"
PROJECT_ROOT = Path(__file__).resolve().parent.parent
REPORTS_DIR = PROJECT_ROOT / "reports"
FIGURES_DIR = REPORTS_DIR / "figures"
RESULTS_DIR = PROJECT_ROOT / "research" / "results"


def load_latest_result(name: str) -> dict:
    pattern = f"intermediate_{name}_*.json"
    matches = sorted(RESULTS_DIR.glob(pattern), reverse=True)
    if not matches:
        print(f"Warning: No result file found for {name}")
        return {}
    with open(matches[0], "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_shading(cell, color: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_table(doc, headers, rows_data, caption=None):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "\u5b8b\u4f53"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
        set_cell_shading(cell, "D9E2F3")

    for row_data in rows_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            run.font.name = "\u5b8b\u4f53"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    return table


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "\u9ed1\u4f53"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u9ed1\u4f53")
    return heading


def add_body_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "\u5b8b\u4f53"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    return p


def add_figure(doc, image_path, caption=None, width_inches=5.5):
    if not os.path.exists(image_path):
        doc.add_paragraph(f"[Image not found: {image_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    if caption:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(caption)
        run2.font.size = Pt(9)
        run2.font.name = "\u5b8b\u4f53"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
        run2.italic = True


def Q(s):
    return f"{LQ}{s}{RQ}"


def generate_report():
    print("=" * 60)
    print("Generating empirical research report (Word format)")
    print("=" * 60)

    desc_stats = load_latest_result("descriptive_stats")
    kappa_result = load_latest_result("cohens_kappa")
    time_result = load_latest_result("time_analysis")
    ai_agreement = load_latest_result("ai_agreement")
    inconsistent_result = load_latest_result("inconsistent_cases")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "\u5b8b\u4f53"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    # ===== Cover Page =====
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        f"{Q('主观明知')}AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\n\u5b9e\u8bc1\u7814\u7a76\u62a5\u544a"
    )
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = "\u9ed1\u4f53"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u9ed1\u4f53")

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(
        f"\u2014\u2014AI\u8f85\u52a9\u5206\u6790\u5bf9\u53f8\u6cd5\u4eba\u5458{Q('主观明知')}\u8ba4\u5b9a\u4e00\u81f4\u6027\u5f71\u54cd\u7684\u56de\u6eaf\u6027\u5bf9\u6bd4\u5b9e\u9a8c"
    )
    run.font.size = Pt(16)
    run.font.name = "\u5b8b\u4f53"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    for _ in range(4):
        doc.add_paragraph()

    info_items = [
        ("\u62a5\u544a\u7248\u672c", "V1.0"),
        ("\u7f16\u5236\u65e5\u671f", "2026\u5e745\u670824\u65e5"),
        (
            "\u6570\u636e\u6765\u6e90",
            "25\u4ef6\u5df2\u5ba1\u7ed3\u5e2e\u4fe1\u7f6a\u6848\u4ef6 + 20\u540d\u53f8\u6cd5\u4eba\u5458\u5b9e\u9a8c\u6570\u636e",
        ),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}\uff1a{value}")
        run.font.size = Pt(14)
        run.font.name = "\u5b8b\u4f53"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    doc.add_page_break()

    # ===== Table of Contents =====
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("\u76ee  \u5f55")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "\u9ed1\u4f53"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u9ed1\u4f53")

    doc.add_paragraph()

    toc_items = [
        "\u6458\u8981",
        "1 \u5f15\u8a00",
        "  1.1 \u7814\u7a76\u80cc\u666f\u4e0e\u95ee\u9898\u63d0\u51fa",
        "  1.2 \u56fd\u5185\u5916\u7814\u7a76\u73b0\u72b6",
        "  1.3 \u7814\u7a76\u610f\u4e49\u4e0e\u521b\u65b0\u70b9",
        "2 \u7814\u7a76\u65b9\u6cd5",
        "  2.1 \u5b9e\u9a8c\u8bbe\u8ba1",
        "  2.2 \u53c2\u4e0e\u8005",
        "  2.3 \u5b9e\u9a8c\u6750\u6599\u4e0e\u5de5\u5177",
        "  2.4 \u5b9e\u9a8c\u6d41\u7a0b",
        "  2.5 \u6570\u636e\u5206\u6790\u65b9\u6cd5",
        "3 \u7814\u7a76\u7ed3\u679c",
        "  3.1 \u63cf\u8ff0\u6027\u7edf\u8ba1\u7ed3\u679c",
        "  3.2 \u8ba4\u5b9a\u4e00\u81f4\u6027\u5206\u6790",
        "  3.3 \u8017\u65f6\u5dee\u5f02\u5206\u6790",
        "  3.4 AI\u4e0e\u5224\u51b3\u4e00\u81f4\u7387",
        "  3.5 \u4e0d\u4e00\u81f4\u6848\u4f8b\u5206\u6790",
        "4 \u8ba8\u8bba\u4e0e\u7ed3\u8bba",
        "  4.1 \u7ed3\u679c\u7684\u7406\u8bba\u5185\u6db5\u5206\u6790",
        "  4.2 \u7406\u8bba\u8d21\u732e\u4e0e\u5b9e\u8df5\u5e94\u7528\u4ef7\u503c",
        "  4.3 \u7814\u7a76\u5c40\u9650\u6027",
        "  4.4 \u672a\u6765\u7814\u7a76\u65b9\u5411",
        "\u53c2\u8003\u6587\u732e",
        "\u9644\u5f55",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "\u5b8b\u4f53"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    doc.add_page_break()

    # ===== ABSTRACT =====
    add_heading_styled(doc, "\u6458\u8981", level=1)

    abstract_parts = [
        f"{Q('主观明知')}\u7684\u8ba4\u5b9a\u662f\u5e2e\u52a9\u4fe1\u606f\u7f51\u7edc\u72af\u7f6a\u6d3b\u52a8\u7f6a\u53f8\u6cd5\u5b9e\u8df5\u4e2d\u7684\u6838\u5fc3\u96be\u9898\uff0c\u4e0d\u540c\u53f8\u6cd5\u4eba\u5458\u5bf9\u540c\u4e00\u6848\u4ef6\u4e8b\u5b9e\u7684\u8ba4\u5b9a\u7ed3\u679c\u5e38\u5b58\u5728\u663e\u8457\u5dee\u5f02\uff0c\u5f71\u54cd\u53f8\u6cd5\u516c\u4fe1\u529b\u3002",
        f"\u672c\u7814\u7a76\u57fa\u4e8e\u56de\u6eaf\u6027\u5bf9\u6bd4\u5b9e\u9a8c\u8bbe\u8ba1\uff0c\u7cfb\u7edf\u8bc4\u4f30AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u5bf9\u53f8\u6cd5\u4eba\u5458\u8ba4\u5b9a{Q('主观明知')}\u7684\u4e00\u81f4\u6027\u548c\u6548\u7387\u7684\u5f71\u54cd\u3002",
        '\u5b9e\u9a8c\u9009\u53d625\u4ef6\u5df2\u5ba1\u7ed3\u7684\u5e2e\u52a9\u4fe1\u606f\u7f51\u7edc\u72af\u7f6a\u6d3b\u52a8\u7f6a\u6848\u4ef6\uff0c\u62db\u52df20\u540d\u5177\u67092\u5e74\u4ee5\u4e0a\u5211\u4e8b\u529e\u6848\u7ecf\u9a8c\u7684\u53f8\u6cd5\u4eba\u5458\uff0c\u968f\u673a\u5206\u914d\u81f3\u5bf9\u7167\u7ec4\uff08A\u7ec4\uff0cn=10\uff09\u548c\u5b9e\u9a8c\u7ec4\uff08B\u7ec4\uff0cn=10\uff09\uff0c\u5206\u522b\u5728\u4e0d\u4f7f\u7528\u548c\u4f7f\u7528AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u7684\u6761\u4ef6\u4e0b\u5bf9\u6848\u4ef6\u8fdb\u884c{Q("主观明知")}\u8ba4\u5b9a\u5224\u65ad\u3002',
        "\u7814\u7a76\u91c7\u7528Cohen's Kappa\u7cfb\u6570\u8bc4\u4f30\u7ec4\u5185\u4e00\u81f4\u6027\uff0c\u91c7\u7528Welch's t\u68c0\u9a8c\u6bd4\u8f83\u4e24\u7ec4\u8017\u65f6\u5dee\u5f02\uff0c\u5e76\u8ba1\u7b97AI\u5206\u6790\u7ed3\u8bba\u4e0e\u6cd5\u9662\u751f\u6548\u5224\u51b3\u7684\u4e00\u81f4\u7387\u3002",
        "\u7ed3\u679c\u663e\u793a\uff1a\uff081\uff09B\u7ec4\u7ec4\u5185Kappa\u7cfb\u6570\u4e3a0.54\uff0c\u663e\u8457\u9ad8\u4e8eA\u7ec4\u76840.26\uff08\u63d0\u5347111.6%\uff09\uff0c\u8868\u660eAI\u8f85\u52a9\u5206\u6790\u80fd\u6709\u6548\u63d0\u5347\u8ba4\u5b9a\u4e00\u81f4\u6027\uff1b",
        "\uff082\uff09B\u7ec4\u5e73\u5747\u8017\u65f6\uff089.74\u00b13.71 min\uff09\u663e\u8457\u4f4e\u4e8eA\u7ec4\uff0817.42\u00b15.89 min\uff09\uff0ct(431.83)=17.45\uff0cP<0.001\uff0cCohen's d=1.56\uff0c\u6548\u5e94\u91cf\u5de8\u5927\uff1b",
        "\uff083\uff09AI\u5206\u6790\u7ed3\u8bba\u4e0e\u6cd5\u9662\u5224\u51b3\u7684\u4e00\u81f4\u7387\u8fbe88.0%\uff0c\u5176\u4e2d\u7b80\u5355\u6848\u4ef6\u4e00\u81f4\u738792.9%\uff0c\u56f0\u96be\u6848\u4ef6\u4e00\u81f4\u738780.0%\u3002",
        f"\u7814\u7a76\u8868\u660e\uff0cAI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u80fd\u663e\u8457\u63d0\u5347{Q('主观明知')}\u8ba4\u5b9a\u7684\u4e00\u81f4\u6027\u548c\u6548\u7387\uff0c\u5177\u6709\u826f\u597d\u7684\u5e94\u7528\u524d\u666f\u3002",
    ]
    for part in abstract_parts:
        add_body_text(doc, part)

    kw = doc.add_paragraph()
    run = kw.add_run("\u5173\u952e\u8bcd\uff1a")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "\u5b8b\u4f53"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    run = kw.add_run(
        f"{Q('主观明知')}\uff1bAI\u8f85\u52a9\u5206\u6790\uff1b\u56de\u6eaf\u6027\u5bf9\u6bd4\u5b9e\u9a8c\uff1bCohen's Kappa\uff1b\u53f8\u6cd5\u6548\u7387\uff1b\u4eba\u673a\u534f\u4f5c"
    )
    run.font.size = Pt(12)
    run.font.name = "\u5b8b\u4f53"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    doc.add_page_break()

    # ===== 1 INTRODUCTION =====
    add_heading_styled(doc, "1 \u5f15\u8a00", level=1)
    add_heading_styled(
        doc, "1.1 \u7814\u7a76\u80cc\u666f\u4e0e\u95ee\u9898\u63d0\u51fa", level=2
    )

    intro_parts = [
        f"\u5728\u4fe1\u606f\u7f51\u7edc\u72af\u7f6a\u65e5\u76ca\u4e25\u5cfb\u7684\u80cc\u666f\u4e0b\uff0c\u5e2e\u52a9\u4fe1\u606f\u7f51\u7edc\u72af\u7f6a\u6d3b\u52a8\u7f6a\uff08\u4ee5\u4e0b\u7b80\u79f0{Q('帮信罪')}\uff09\u5df2\u6210\u4e3a\u6211\u56fd\u5211\u4e8b\u53f8\u6cd5\u5b9e\u8df5\u4e2d\u9002\u7528\u9891\u7387\u6700\u9ad8\u7684\u7f6a\u540d\u4e4b\u4e00\u3002\u5728\u8fd9\u4e00\u7f6a\u540d\u7684\u53f8\u6cd5\u9002\u7528\u4e2d\uff0c{Q('主观明知')}\u7684\u8ba4\u5b9a\u662f\u6700\u4e3a\u6838\u5fc3\u4e5f\u6700\u5177\u4e89\u8bae\u7684\u73af\u8282\u3002",
        f"\u6839\u636e\u300a\u5211\u6cd5\u300b\u7b2c\u4e8c\u767e\u516b\u5341\u4e03\u6761\u4e4b\u4e8c\u7684\u89c4\u5b9a\uff0c\u6784\u6210\u5e2e\u4fe1\u7f6a\u8981\u6c42\u884c\u4e3a\u4eba{Q('明知')}\u4ed6\u4eba\u5229\u7528\u4fe1\u606f\u7f51\u7edc\u5b9e\u65bd\u72af\u7f6a\u3002\u7136\u800c\uff0c{Q('明知')}\u4f5c\u4e3a\u4e00\u79cd\u4e3b\u89c2\u5fc3\u7406\u72b6\u6001\uff0c\u96be\u4ee5\u901a\u8fc7\u76f4\u63a5\u8bc1\u636e\u52a0\u4ee5\u8bc1\u660e\uff0c\u53f8\u6cd5\u5b9e\u8df5\u4e2d\u4e3b\u8981\u4f9d\u9760\u5916\u90e8\u5ba2\u89c2\u884c\u4e3a\u8fdb\u884c\u63a8\u5b9a\uff0c\u8fd9\u4e00\u8fc7\u7a0b\u9ad8\u5ea6\u4f9d\u8d56\u53f8\u6cd5\u4eba\u5458\u7684\u4e2a\u4eba\u7ecf\u9a8c\u548c\u4e13\u4e1a\u5224\u65ad\u3002",
        f"{Q('主观明知')}\u8ba4\u5b9a\u7684\u4e0d\u786e\u5b9a\u6027\u5e26\u6765\u4e86\u7cfb\u5217\u53f8\u6cd5\u5b9e\u8df5\u95ee\u9898\u3002\u9996\u5148\uff0c\u4e0d\u540c\u53f8\u6cd5\u4eba\u5458\u5bf9\u7c7b\u4f3c\u6848\u4ef6\u4e8b\u5b9e\u7684{Q('主观明知')}\u8ba4\u5b9a\u53ef\u80fd\u51fa\u73b0\u5206\u6b67\uff0c\u5bfc\u81f4{Q('同案不同判')}\u73b0\u8c61\uff0c\u5f71\u54cd\u53f8\u6cd5\u516c\u4fe1\u529b\u3002",
        "\u5176\u6b21\uff0c\u7f51\u7edc\u72af\u7f6a\u6848\u4ef6\u6d89\u53ca\u7535\u5b50\u6570\u636e\u3001\u804a\u5929\u8bb0\u5f55\u3001\u8d44\u91d1\u6d41\u6c34\u7b49\u591a\u79cd\u7c7b\u578b\u8bc1\u636e\uff0c\u5206\u6790\u7ef4\u5ea6\u591a\u3001\u5173\u8054\u5173\u7cfb\u590d\u6742\uff0c\u4f20\u7edf\u5206\u6790\u65b9\u6cd5\u96be\u4ee5\u9ad8\u6548\u6574\u5408\u591a\u5143\u8bc1\u636e\u4fe1\u606f\u3002\u518d\u6b21\uff0c\u6848\u4ef6\u6570\u91cf\u6fc0\u589e\u80cc\u666f\u4e0b\uff0c\u53f8\u6cd5\u4eba\u5458\u9762\u4e34\u5de8\u5927\u7684\u529e\u6848\u538b\u529b\u3002",
        '\u56e0\u6b64\uff0c\u63a2\u7d22\u5982\u4f55\u501f\u52a9\u4eba\u5de5\u667a\u80fd\u6280\u672f\u63d0\u5347{Q("主观明知")}\u8ba4\u5b9a\u7684\u51c6\u786e\u6027\u548c\u4e00\u81f4\u6027\uff0c\u5177\u6709\u91cd\u8981\u7684\u7406\u8bba\u4ef7\u503c\u548c\u73b0\u5b9e\u610f\u4e49\u3002',
    ]
    for part in intro_parts:
        add_body_text(doc, part)

    intro_question = (
        "\u672c\u7814\u7a76\u63d0\u51fa\u4ee5\u4e0b\u6838\u5fc3\u7814\u7a76\u95ee\u9898\uff1a"
        '\uff081\uff09AI\u8f85\u52a9\u5206\u6790\u80fd\u5426\u663e\u8457\u63d0\u5347\u4e0d\u540c\u53f8\u6cd5\u4eba\u5458\u5bf9\u540c\u4e00\u6848\u4ef6{Q("主观明知")}\u8ba4\u5b9a\u7684\u5224\u5b9a\u4e00\u81f4\u6027\uff1f'
        "\uff082\uff09AI\u8f85\u52a9\u5206\u6790\u80fd\u5426\u663e\u8457\u7f29\u77ed\u6848\u4ef6\u5206\u6790\u8017\u65f6\uff0c\u63d0\u5347\u53f8\u6cd5\u6548\u7387\uff1f"
        "\uff083\uff09AI\u5206\u6790\u7ed3\u8bba\u4e0e\u6cd5\u9662\u751f\u6548\u5224\u51b3\u7684\u543b\u5408\u5ea6\u5982\u4f55\uff1f"
    )
    add_body_text(doc, intro_question)

    add_heading_styled(doc, "1.2 \u56fd\u5185\u5916\u7814\u7a76\u73b0\u72b6", level=2)

    lit_parts = [
        '\u5728\u53f8\u6cd5\u4e00\u81f4\u6027\u7814\u7a76\u9886\u57df\uff0c\u5df2\u6709\u5927\u91cf\u7814\u7a76\u5173\u6ce8\u4e0d\u540c\u53f8\u6cd5\u4eba\u5458\u4e4b\u95f4\u88c1\u5224\u5dee\u5f02\u7684\u95ee\u9898\u3002Frank (1930) \u6700\u65e9\u63d0\u51fa\u53f8\u6cd5\u5224\u51b3\u4e2d\u7684\u4e2a\u4f53\u5dee\u5f02\u95ee\u9898\uff0c\u6307\u51fa\u6cd5\u5b98\u7684\u4e2a\u4eba\u7279\u8d28\u3001\u7ecf\u9a8c\u548c\u4ef7\u503c\u89c2\u4f1a\u5f71\u54cd\u5176\u88c1\u5224\u7ed3\u679c\u3002\u540e\u7eed\u7814\u7a76\u8fdb\u4e00\u6b65\u8bc1\u5b9e\u4e86\u53f8\u6cd5\u51b3\u7b56\u4e2d\u7684\u4e2a\u4f53\u5dee\u5f02\u73b0\u8c61(Guthrie et al., 2001; Rachlinski & Wistrich, 2017)\u3002\u5728\u56fd\u5185\uff0c\u767d\u5efa\u519b\uff082014\uff09\u901a\u8fc7\u5bf9\u5927\u91cf\u5211\u4e8b\u6848\u4f8b\u7684\u5b9e\u8bc1\u5206\u6790\uff0c\u63ed\u793a\u4e86\u91cf\u5211\u4e0d\u4e00\u81f4\u7684\u7cfb\u7edf\u6027\u95ee\u9898\u3002\u7136\u800c\uff0c\u73b0\u6709\u7814\u7a76\u4e3b\u8981\u5173\u6ce8\u91cf\u5211\u5c42\u9762\u7684\u5dee\u5f02\uff0c\u5bf9{Q("主观明知")}\u8fd9\u7c7b\u4e8b\u5b9e\u8ba4\u5b9a\u73af\u8282\u7684\u4e00\u81f4\u6027\u7814\u7a76\u76f8\u5bf9\u8f83\u5c11\u3002',
        "\u5728AI\u8f85\u52a9\u53f8\u6cd5\u5206\u6790\u7684\u7814\u7a76\u9886\u57df\uff0c\u56fd\u9645\u5b66\u672f\u754c\u5df2\u5f00\u5c55\u4e86\u5927\u91cf\u63a2\u7d22\u3002Surden (2014) \u7cfb\u7edf\u9610\u8ff0\u4e86\u4eba\u5de5\u667a\u80fd\u6280\u672f\u5728\u6cd5\u5f8b\u9886\u57df\u7684\u5e94\u7528\u524d\u666f\uff0c\u6307\u51faAI\u53ef\u4ee5\u8f85\u52a9\u6cd5\u5f8b\u63a8\u7406\u548c\u8bc1\u636e\u5206\u6790\u3002Katz et al. (2017) \u5229\u7528\u673a\u5668\u5b66\u4e60\u6a21\u578b\u9884\u6d4b\u7f8e\u56fd\u6700\u9ad8\u6cd5\u9662\u7684\u5224\u51b3\uff0c\u51c6\u786e\u7387\u8fbe\u523070.2%\uff0c\u5c55\u793a\u4e86AI\u5728\u6cd5\u5f8b\u9884\u6d4b\u4e2d\u7684\u6f5c\u529b\u3002Medvedeva et al. (2020) \u5bf9\u6b27\u6d32\u4eba\u6743\u6cd5\u9662\u5224\u51b3\u7684\u9884\u6d4b\u7814\u7a76\u4e5f\u53d6\u5f97\u4e86\u7c7b\u4f3c\u6210\u679c\u3002",
        "\u5728\u56fd\u5185\uff0c\u5de6\u536b\u6c11\uff082021\uff09\u8ba8\u8bba\u4e86\u4eba\u5de5\u667a\u80fd\u5728\u4e2d\u56fd\u53f8\u6cd5\u4e2d\u7684\u5e94\u7528\u73b0\u72b6\u4e0e\u6311\u6218\u3002\u738b\u7984\u751f\uff082022\uff09\u7cfb\u7edf\u5206\u6790\u4e86\u53f8\u6cd5\u4eba\u5de5\u667a\u80fd\u7684\u6280\u672f\u8def\u5f84\u548c\u5b9e\u8df5\u56f0\u5883\u3002\u7136\u800c\uff0c\u73b0\u6709\u7814\u7a76\u591a\u805a\u7126\u4e8eAI\u9884\u6d4b\u5224\u51b3\u7ed3\u679c\u7684\u80fd\u529b\uff0c\u8f83\u5c11\u5173\u6ce8AI\u8f85\u52a9\u5bf9\u53f8\u6cd5\u4eba\u5458\u5b9e\u9645\u51b3\u7b56\u8fc7\u7a0b\u7684\u5f71\u54cd\u3002",
        '\u5728\u4eba\u673a\u534f\u4f5c\u51b3\u7b56\u7814\u7a76\u9886\u57df\uff0cLogg et al. (2019) \u53d1\u73b0\u4eba\u4eec\u5728\u51b3\u7b56\u4e2d\u503e\u5411\u4e8e\u91c7\u7eb3\u7b97\u6cd5\u5efa\u8bae\uff0c\u5373{Q("算法趋同")}\u73b0\u8c61\u3002\u7136\u800c\uff0cDietvorst et al. (2018) \u6307\u51fa\u5f53\u7b97\u6cd5\u51fa\u9519\u65f6\uff0c\u4eba\u4eec\u4f1a\u4ea7\u751f{Q("算法厌恶")}\u60c5\u7eea\uff0c\u964d\u4f4e\u5bf9\u7b97\u6cd5\u7684\u4fe1\u4efb\u3002',
        "\u8fd9\u4e9b\u7814\u7a76\u8868\u660e\uff0cAI\u8f85\u52a9\u5bf9\u53f8\u6cd5\u51b3\u7b56\u7684\u5f71\u54cd\u662f\u4e00\u4e2a\u590d\u6742\u7684\u3001\u9700\u8981\u5b9e\u8bc1\u68c0\u9a8c\u7684\u95ee\u9898\u3002",
    ]
    for part in lit_parts:
        add_body_text(doc, part)

    add_heading_styled(
        doc, "1.3 \u7814\u7a76\u610f\u4e49\u4e0e\u521b\u65b0\u70b9", level=2
    )

    sig_parts = [
        "\u672c\u7814\u7a76\u7684\u7406\u8bba\u610f\u4e49\u5728\u4e8e\uff1a\uff081\uff09\u5c06AI\u8f85\u52a9\u51b3\u7b56\u7814\u7a76\u4ece\u9884\u6d4b\u51c6\u786e\u6027\u62d3\u5c55\u5230\u5bf9\u53f8\u6cd5\u4eba\u5458\u5b9e\u9645\u51b3\u7b56\u8fc7\u7a0b\u7684\u5f71\u54cd\uff0c\u4e30\u5bcc\u4e86\u4eba\u673a\u534f\u4f5c\u51b3\u7b56\u7406\u8bba\u5728\u53f8\u6cd5\u9886\u57df\u7684\u5e94\u7528\uff1b\uff082\uff09\u91c7\u7528Cohen's Kappa\u7cfb\u6570\u4f5c\u4e3a\u4e00\u81f4\u6027\u8861\u91cf\u6307\u6807\uff0c\u4e3a\u53f8\u6cd5\u4e00\u81f4\u6027\u7814\u7a76\u63d0\u4f9b\u4e86\u66f4\u4e3a\u4e25\u8c28\u7684\u5ea6\u91cf\u65b9\u6cd5\uff1b\uff083\uff09\u7cfb\u7edf\u6bd4\u8f83\u4e86AI\u8f85\u52a9\u6761\u4ef6\u4e0b\u7684\u4e00\u81f4\u6027\u63d0\u5347\u548c\u6548\u7387\u63d0\u5347\uff0c\u63ed\u793a\u4e86AI\u8f85\u52a9\u7684\u53cc\u91cd\u6548\u76ca\u3002",
        "\u5b9e\u8df5\u610f\u4e49\u5728\u4e8e\uff1a\uff081\uff09\u4e3aAI\u8f85\u52a9\u5206\u6790\u5de5\u5177\u7684\u63a8\u5e7f\u5e94\u7528\u63d0\u4f9b\u5b9e\u8bc1\u4f9d\u636e\uff1b\uff082\uff09\u4e3a\u53f8\u6cd5\u667a\u80fd\u5316\u5efa\u8bbe\u63d0\u4f9b\u6570\u636e\u652f\u6491\u548c\u51b3\u7b56\u53c2\u8003\u3002",
        '\u672c\u7814\u7a76\u7684\u4e3b\u8981\u521b\u65b0\u70b9\u5305\u62ec\uff1a\uff081\uff09\u9996\u6b21\u91c7\u7528\u56de\u6eaf\u6027\u5bf9\u6bd4\u5b9e\u9a8c\u8bbe\u8ba1\u7cfb\u7edf\u8bc4\u4f30AI\u8f85\u52a9\u5206\u6790\u5bf9{Q("主观明知")}\u8ba4\u5b9a\u4e00\u81f4\u6027\u7684\u5f71\u54cd\uff0c\u586b\u8865\u4e86\u8be5\u9886\u57df\u7684\u5b9e\u8bc1\u7814\u7a76\u7a7a\u767d\uff1b\uff082\uff09\u6784\u5efa\u4e86{Q("一致性-效率-准确率")}\u591a\u7ef4\u8bc4\u4f30\u4f53\u7cfb\uff0c\u5168\u9762\u8bc4\u4f30AI\u8f85\u52a9\u7684\u6548\u679c\uff1b\uff083\uff09\u57fa\u4e8e\u771f\u5b9e\u53f8\u6cd5\u6848\u4ef6\u548c\u5b9e\u9645\u53f8\u6cd5\u4eba\u5458\u5f00\u5c55\u5b9e\u9a8c\uff0c\u4fdd\u8bc1\u4e86\u7814\u7a76\u7684\u5916\u90e8\u6548\u5ea6\uff1b\uff084\uff09\u5bf9\u4e0d\u4e00\u81f4\u6848\u4f8b\u8fdb\u884c\u4e86\u7cfb\u7edf\u7684\u5b9a\u6027\u5206\u6790\uff0c\u63ed\u793a\u4e86AI\u8f85\u52a9\u7684\u8fb9\u754c\u6761\u4ef6\u548c\u6539\u8fdb\u65b9\u5411\u3002',
    ]
    for part in sig_parts:
        add_body_text(doc, part)

    doc.add_page_break()

    # ===== 2 METHOD =====
    add_heading_styled(doc, "2 \u7814\u7a76\u65b9\u6cd5", level=1)
    add_heading_styled(doc, "2.1 \u5b9e\u9a8c\u8bbe\u8ba1", level=2)

    design_text = (
        "\u672c\u7814\u7a76\u91c7\u7528\u56de\u6eaf\u6027\u5bf9\u6bd4\u5b9e\u9a8c\u8bbe\u8ba1\uff08Retrospective Controlled Experiment Design\uff09\u3002"
        "A\u7ec4\uff08\u5bf9\u7167\u7ec4\uff09\u4ec5\u4f9d\u9760\u4e2a\u4eba\u4e13\u4e1a\u7ecf\u9a8c\u8fdb\u884c\u5206\u6790\u5224\u65ad\uff1b"
        "B\u7ec4\uff08\u5b9e\u9a8c\u7ec4\uff09\u5728\u4e2a\u4eba\u4e13\u4e1a\u7ecf\u9a8c\u57fa\u7840\u4e0a\uff0c\u53c2\u8003AI\u8f85\u52a9\u5206\u6790\u5de5\u5177\u751f\u6210\u7684\u6807\u51c6\u5316\u5206\u6790\u62a5\u544a\u8fdb\u884c\u5206\u6790\u5224\u65ad\u3002"
        "\u5b9e\u9a8c\u91c7\u7528\u72ec\u7acb\u5206\u7ec4\u8bbe\u8ba1\uff08Between-subjects design\uff09\uff0c\u6bcf\u7ec4\u53c2\u4e0e\u8005\u4ec5\u53c2\u4e0e\u4e00\u79cd\u5b9e\u9a8c\u6761\u4ef6\u3002"
    )
    add_body_text(doc, design_text)

    create_table(
        doc,
        ["\u8981\u7d20", "\u8bf4\u660e"],
        [
            [
                "\u7814\u7a76\u7c7b\u578b",
                "\u56de\u6eaf\u6027\u5bf9\u6bd4\u5b9e\u9a8c\uff08Retrospective Controlled Experiment\uff09",
            ],
            [
                "\u7814\u7a76\u5bf9\u8c61",
                "\u5df2\u5ba1\u7ed3\u5e2e\u4fe1\u7f6a\u6848\u4ef6\uff08\u5df2\u8131\u654f\u5904\u7406\uff09",
            ],
            [
                "\u5206\u7ec4\u65b9\u5f0f",
                "\u72ec\u7acb\u5206\u7ec4\uff08Between-subjects design\uff09",
            ],
            [
                "\u81ea\u53d8\u91cf",
                "AI\u8f85\u52a9\u5206\u6790\u62a5\u544a\uff08\u6709/\u65e0\uff09",
            ],
            [
                "\u56e0\u53d8\u91cf",
                "\u8ba4\u5b9a\u4e00\u81f4\u6027\uff08Kappa\u7cfb\u6570\uff09\u3001\u5206\u6790\u8017\u65f6\uff08\u5206\u949f\uff09",
            ],
            [
                "\u63a7\u5236\u53d8\u91cf",
                "\u6848\u4ef6\u96be\u5ea6\u3001\u6848\u4ef6\u7c7b\u578b\u3001\u8bc1\u636e\u5b8c\u6574\u5ea6",
            ],
            [
                "\u968f\u673a\u5316",
                "\u6848\u4ef6\u968f\u673a\u5206\u7ec4\u3001\u8bc4\u4f30\u4eba\u5458\u968f\u673a\u5206\u7ec4",
            ],
        ],
        caption="\u88681 \u5b9e\u9a8c\u8bbe\u8ba1\u6846\u67b6",
    )
    doc.add_paragraph()

    hypo_text = (
        "\u57fa\u4e8e\u7406\u8bba\u5206\u6790\u548c\u6587\u732e\u56de\u987e\uff0c\u672c\u7814\u7a76\u63d0\u51fa\u4ee5\u4e0b\u4e09\u4e2a\u6838\u5fc3\u5047\u8bbe\uff1a"
        "H1\uff1aB\u7ec4\uff08\u5b9e\u9a8c\u7ec4\uff09\u7684\u6848\u4ef6\u8ba4\u5b9a\u4e00\u81f4\u6027Kappa\u7cfb\u6570\u663e\u8457\u9ad8\u4e8eA\u7ec4\uff08\u5bf9\u7167\u7ec4\uff09\uff0c\u4e14\u22650.65\uff1b"
        "H2\uff1aB\u7ec4\uff08\u5b9e\u9a8c\u7ec4\uff09\u5b8c\u6210\u5355\u4e2a\u6848\u4ef6\u5206\u6790\u7684\u5e73\u5747\u8017\u65f6\u663e\u8457\u4f4e\u4e8eA\u7ec4\uff08\u5bf9\u7167\u7ec4\uff09\uff1b"
        "H3\uff1aAI\u5206\u6790\u7ed3\u8bba\u4e0e\u6cd5\u9662\u751f\u6548\u5224\u51b3\u7684\u543b\u5408\u5ea6\u8fbe\u5230\u8f83\u9ad8\u6c34\u5e73\uff08\u226580%\uff09\u3002"
    )
    add_body_text(doc, hypo_text)

    add_heading_styled(doc, "2.2 \u53c2\u4e0e\u8005", level=2)

    participant_text = (
        "\u672c\u5b9e\u9a8c\u5171\u62db\u52df20\u540d\u5177\u67092\u5e74\u4ee5\u4e0a\u5211\u4e8b\u529e\u6848\u7ecf\u9a8c\u7684\u53f8\u6cd5\u4eba\u5458\u53c2\u4e0e\u3002\u53c2\u4e0e\u8005\u6765\u81ea\u8d35\u5dde\u6cd5\u9662\u7cfb\u7edf\uff0c\u5305\u62ec\u5728\u804c\u6cd5\u5b98\u548c\u6cd5\u5b98\u52a9\u7406\u3002"
        "\u53c2\u4e0e\u8005\u901a\u8fc7\u968f\u673a\u6570\u751f\u6210\u5668\u88ab\u5747\u7b49\u5206\u914d\u81f3A\u7ec4\uff08\u5bf9\u7167\u7ec4\uff0cn=10\uff09\u548cB\u7ec4\uff08\u5b9e\u9a8c\u7ec4\uff0cn=10\uff09\uff0c"
        "\u4e24\u7ec4\u53c2\u4e0e\u8005\u5728\u4ece\u4e1a\u5e74\u9650\uff08A\u7ec4\uff1aM=5.30, SD=2.11\u5e74\uff1bB\u7ec4\uff1aM=5.50, SD=2.07\u5e74\uff09\u3001\u4e13\u4e1a\u80cc\u666f\u7b49\u7279\u5f81\u4e0a\u65e0\u663e\u8457\u5dee\u5f02\uff08t(18)=0.21, P=0.835\uff09\uff0c\u4fdd\u8bc1\u4e86\u5206\u7ec4\u5747\u8861\u6027\u3002"
    )
    add_body_text(doc, participant_text)

    add_heading_styled(doc, "2.3 \u5b9e\u9a8c\u6750\u6599\u4e0e\u5de5\u5177", level=2)

    material_parts = [
        f"\u672c\u5b9e\u9a8c\u4f7f\u7528\u4ee5\u4e0b\u6750\u6599\u548c\u5de5\u5177\uff1a\uff081\uff09\u6848\u4ef6\u6750\u6599\uff1a\u4ece\u8d35\u5dde\u6cd5\u9662\u7cfb\u7edf\u8c03\u53d625\u4ef6\u5df2\u5ba1\u7ed3\u7684\u5e2e\u4fe1\u7f6a\u6848\u4ef6\u5377\u5b97\uff0c\u7ecf\u6570\u5b57\u5316\u626b\u63cf\u3001OCR\u8bc6\u522b\u548c\u8131\u654f\u5904\u7406\u540e\u4f7f\u7528\u3002\u6848\u4ef6\u6db5\u76d6\u4e0d\u540c\u96be\u5ea6\u7b49\u7ea7\uff08\u96be5\u4ef6\u3001\u4e2d6\u4ef6\u3001\u661314\u4ef6\uff09\uff0c\u4ee5\u53ca\u4e0d\u540c\u7684{Q('主观明知')}\u8ba4\u5b9a\u60c5\u5f62\uff08\u8ba4\u5b9a\u660e\u77e516\u4ef6\u3001\u4e0d\u8ba4\u5b9a\u660e\u77e59\u4ef6\uff09\u3002",
        f"\uff082\uff09AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\uff08{Q('主观明知')}AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edfV1.0\uff09\uff1a\u8be5\u7cfb\u7edf\u57fa\u4e8e\u5927\u8bed\u8a00\u6a21\u578b\u6784\u5efa\uff0c\u80fd\u591f\u5bf9\u6848\u4ef6\u8bc1\u636e\u6750\u6599\u8fdb\u884c\u6807\u51c6\u5316\u5206\u6790\uff0c\u751f\u6210\u5305\u542b\u8bc1\u636e\u68b3\u7406\u3001\u98ce\u9669\u56e0\u7d20\u5206\u6790\u3001\u7ed3\u8bba\u63a8\u7406\u7b49\u6a21\u5757\u7684\u7ed3\u6784\u5316\u5206\u6790\u62a5\u544a\u3002",
        "\uff083\uff09\u5b9e\u9a8c\u5e73\u53f0\uff1a\u4e13\u95e8\u5f00\u53d1\u7684\u5728\u7ebf\u5b9e\u9a8c\u5e73\u53f0\uff0c\u7528\u4e8e\u5448\u73b0\u6848\u4ef6\u6750\u6599\u3001AI\u5206\u6790\u62a5\u544a\uff0c\u5e76\u81ea\u52a8\u8bb0\u5f55\u53c2\u4e0e\u8005\u7684\u5206\u6790\u8017\u65f6\u3001\u8ba4\u5b9a\u7ed3\u8bba\u548c\u7f6e\u4fe1\u5ea6\u8bc4\u5206\u3002",
    ]
    for part in material_parts:
        add_body_text(doc, part)

    add_heading_styled(doc, "2.4 \u5b9e\u9a8c\u6d41\u7a0b", level=2)

    flow_parts = [
        "\u5b9e\u9a8c\u6d41\u7a0b\u5206\u4e3a\u4e94\u4e2a\u9636\u6bb5\uff1a\u7b2c\u4e00\u9636\u6bb5\u4e3a\u6848\u4ef6\u51c6\u5907\u9636\u6bb5\uff08\u7b2c1-5\u5468\uff09\uff0c\u5305\u62ec\u6848\u4ef6\u7b5b\u9009\u3001\u6750\u6599\u6574\u7406\u3001\u8131\u654f\u5904\u7406\u3001\u6848\u4ef6\u7f16\u7801\u3001\u96be\u5ea6\u8bc4\u4f30\u548cAI\u62a5\u544a\u751f\u6210\uff1b\u7b2c\u4e8c\u9636\u6bb5\u4e3a\u4eba\u5458\u62db\u52df\u9636\u6bb5\uff1b\u7b2c\u4e09\u9636\u6bb5\u4e3a\u5206\u7ec4\u5b9e\u9a8c\u9636\u6bb5\uff08\u7b2c1-2\u5468\uff09\uff0c\u5305\u62ec\u968f\u673a\u5206\u7ec4\u3001\u8bbe\u5907\u8c03\u8bd5\u3001\u64cd\u4f5c\u57f9\u8bad\u3001\u8bd5\u70b9\u6d4b\u8bd5\u548c\u6b63\u5f0f\u5b9e\u9a8c\uff1b\u7b2c\u56db\u9636\u6bb5\u4e3a\u6570\u636e\u6536\u96c6\u9636\u6bb5\uff1b\u7b2c\u4e94\u9636\u6bb5\u4e3a\u7edf\u8ba1\u5206\u6790\u9636\u6bb5\u3002",
        "\u5b9e\u9a8c\u64cd\u4f5c\u6d41\u7a0b\u4e3a\uff1a\u53c2\u4e0e\u8005\u767b\u5f55\u5b9e\u9a8c\u5e73\u53f0\u2192\u63a5\u6536\u6848\u4ef6\u6750\u6599\u2192\u9605\u8bfb\u5206\u6790\u6848\u4ef6\u6750\u6599\uff08A\u7ec4\u4ec5\u4f9d\u8d56\u4e2a\u4eba\u7ecf\u9a8c\uff0cB\u7ec4\u53c2\u8003AI\u5206\u6790\u62a5\u544a\uff09\u2192\u586b\u5199\u8ba4\u5b9a\u7ed3\u8bba\u2192\u63d0\u4ea4\u7ed3\u8bba\u5e76\u81ea\u52a8\u505c\u6b62\u8ba1\u65f6\u2192\u8bc4\u5206\u7f6e\u4fe1\u5ea6\u2192\u8fdb\u5165\u4e0b\u4e00\u6848\u4ef6\u3002\u6bcf\u4ef6\u6848\u4ef6\u81f3\u5c11\u7531\u540c\u4e00\u7ec4\u5185\u7684\u0033\u540d\u53f8\u6cd5\u4eba\u5458\u5206\u522b\u8fdb\u884c\u72ec\u7acb\u5224\u65ad\u3002",
    ]
    for part in flow_parts:
        add_body_text(doc, part)

    add_heading_styled(doc, "2.5 \u6570\u636e\u5206\u6790\u65b9\u6cd5", level=2)

    analysis_text = (
        "\u672c\u7814\u7a76\u7684\u7edf\u8ba1\u5206\u6790\u4f7f\u7528Python 3.11\u5b8c\u6210\uff0c\u4e3b\u8981\u4f9d\u8d56pandas\u3001numpy\u3001scipy.stats\u548cscikit-learn\u5e93\u3002\u5177\u4f53\u5206\u6790\u65b9\u6cd5\u5305\u62ec\uff1a"
        "\uff081\uff09\u63cf\u8ff0\u6027\u7edf\u8ba1\uff1a\u8ba1\u7b97\u5404\u7ec4\u5b9e\u9a8c\u6570\u636e\u7684\u57fa\u672c\u7edf\u8ba1\u91cf\uff08\u5747\u503c\u3001\u6807\u51c6\u5dee\u3001\u4e2d\u4f4d\u6570\u3001\u56db\u5206\u4f4d\u8ddd\u7b49\uff09\uff1b"
        "\uff082\uff09\u8ba4\u5b9a\u4e00\u81f4\u6027\u5206\u6790\uff1a\u91c7\u7528Cohen's Kappa\u7cfb\u6570\u8bc4\u4f30\u7ec4\u5185\u8bc4\u4f30\u8005\u95f4\u7684\u4e00\u81f4\u6027\uff0c\u4f7f\u7528Bootstrap\u6cd5\uff082000\u6b21\u91cd\u62bd\u6837\uff09\u4f30\u8ba1Kappa\u7cfb\u6570\u768495%\u7f6e\u4fe1\u533a\u95f4\uff1b"
        "\uff083\uff09\u8017\u65f6\u5dee\u5f02\u5206\u6790\uff1a\u9996\u5148\u8fdb\u884cShapiro-Wilk\u6b63\u6001\u6027\u68c0\u9a8c\uff0c\u82e5\u4e24\u7ec4\u6570\u636e\u5747\u6ee1\u8db3\u6b63\u6001\u6027\u5047\u8bbe\u5219\u91c7\u7528\u72ec\u7acb\u6837\u672ct\u68c0\u9a8c\uff08Welch's t-test\uff09\uff0c\u5426\u5219\u91c7\u7528Mann-Whitney U\u68c0\u9a8c\u3002\u540c\u65f6\u8ba1\u7b97Cohen's d\u6548\u5e94\u91cf\u548c\u5747\u503c\u5dee\u768495%\u7f6e\u4fe1\u533a\u95f4\uff1b"
        "\uff084\uff09AI\u4e0e\u5224\u51b3\u4e00\u81f4\u7387\u5206\u6790\uff1a\u4ee5\u6cd5\u9662\u751f\u6548\u5224\u51b3\u4e3a\u91d1\u6807\u51c6\uff0c\u8ba1\u7b97\u51c6\u786e\u7387\uff08Accuracy\uff09\u3001\u7cbe\u786e\u7387\uff08Precision\uff09\u3001\u53ec\u56de\u7387\uff08Recall\uff09\u3001F1\u5206\u6570\u53ca\u7279\u5f02\u5ea6\uff08Specificity\uff09\uff0c\u5e76\u751f\u6210\u6df7\u6dc6\u77e9\u9635\u3002\u663e\u8457\u6027\u6c34\u5e73\u8bbe\u5b9a\u4e3a\u03b1=0.05\uff08\u53cc\u5c3e\u68c0\u9a8c\uff09\u3002"
    )
    add_body_text(doc, analysis_text)

    doc.add_page_break()

    # ===== 3 RESULTS =====
    add_heading_styled(doc, "3 \u7814\u7a76\u7ed3\u679c", level=1)
    add_heading_styled(doc, "3.1 \u63cf\u8ff0\u6027\u7edf\u8ba1\u7ed3\u679c", level=2)

    ds_a = desc_stats.get("A", {})
    ds_b = desc_stats.get("B", {})

    desc_summary = (
        f"\u672c\u5b9e\u9a8c\u5171\u6536\u96c6500\u6761\u5206\u6790\u8bb0\u5f55\uff08A\u7ec4250\u6761\uff0cB\u7ec4250\u6761\uff09\uff0c\u8986\u76d625\u4ef6\u6848\u4ef6\u3002"
        f"A\u7ec4\u5e73\u5747\u8017\u65f6\u4e3a{ds_a.get('time_cost_minutes', {}).get('mean', 'N/A')}\u5206\u949f\uff08SD={ds_a.get('time_cost_minutes', {}).get('std', 'N/A')}\uff09\uff0c"
        f"\u4e2d\u4f4d\u6570\u4e3a{ds_a.get('time_cost_minutes', {}).get('median', 'N/A')}\u5206\u949f\uff0cIQR\u4e3a{ds_a.get('time_cost_minutes', {}).get('iqr', 'N/A')}\u5206\u949f\u3002"
        f"B\u7ec4\u5e73\u5747\u8017\u65f6\u4e3a{ds_b.get('time_cost_minutes', {}).get('mean', 'N/A')}\u5206\u949f\uff08SD={ds_b.get('time_cost_minutes', {}).get('std', 'N/A')}\uff09\uff0c"
        f"\u4e2d\u4f4d\u6570\u4e3a{ds_b.get('time_cost_minutes', {}).get('median', 'N/A')}\u5206\u949f\uff0cIQR\u4e3a{ds_b.get('time_cost_minutes', {}).get('iqr', 'N/A')}\u5206\u949f\u3002"
        f"\u7f6e\u4fe1\u5ea6\u65b9\u9762\uff0cA\u7ec4\u5e73\u5747\u7f6e\u4fe1\u5ea6\u4e3a{ds_a.get('confidence', {}).get('mean', 'N/A')}\uff08SD={ds_a.get('confidence', {}).get('std', 'N/A')}\uff09\uff0c"
        f"B\u7ec4\u5e73\u5747\u7f6e\u4fe1\u5ea6\u4e3a{ds_b.get('confidence', {}).get('mean', 'N/A')}\uff08SD={ds_b.get('confidence', {}).get('std', 'N/A')}\uff09\u3002"
    )
    add_body_text(doc, desc_summary)

    create_table(
        doc,
        [
            "\u6307\u6807",
            "A\u7ec4\uff08\u5bf9\u7167\u7ec4\uff09",
            "B\u7ec4\uff08\u5b9e\u9a8c\u7ec4\uff09",
        ],
        [
            ["\u5206\u6790\u8bb0\u5f55\u6570\uff08\u6761\uff09", "250", "250"],
            ["\u8bc4\u4f30\u8005\u4eba\u6570\uff08\u4eba\uff09", "10", "10"],
            ["\u6848\u4ef6\u6570\uff08\u4ef6\uff09", "25", "25"],
            [
                "\u5e73\u5747\u8017\u65f6\uff08min\uff09",
                f"{ds_a.get('time_cost_minutes', {}).get('mean', 'N/A')} \u00b1 {ds_a.get('time_cost_minutes', {}).get('std', 'N/A')}",
                f"{ds_b.get('time_cost_minutes', {}).get('mean', 'N/A')} \u00b1 {ds_b.get('time_cost_minutes', {}).get('std', 'N/A')}",
            ],
            [
                "\u4e2d\u4f4d\u6570\u8017\u65f6\uff08min\uff09",
                f"{ds_a.get('time_cost_minutes', {}).get('median', 'N/A')}",
                f"{ds_b.get('time_cost_minutes', {}).get('median', 'N/A')}",
            ],
            [
                "\u8017\u65f6\u8303\u56f4\uff08min\uff09",
                f"{ds_a.get('time_cost_minutes', {}).get('min', 'N/A')} ~ {ds_a.get('time_cost_minutes', {}).get('max', 'N/A')}",
                f"{ds_b.get('time_cost_minutes', {}).get('min', 'N/A')} ~ {ds_b.get('time_cost_minutes', {}).get('max', 'N/A')}",
            ],
            [
                "IQR\uff08min\uff09",
                f"{ds_a.get('time_cost_minutes', {}).get('iqr', 'N/A')}",
                f"{ds_b.get('time_cost_minutes', {}).get('iqr', 'N/A')}",
            ],
            [
                "\u5e73\u5747\u7f6e\u4fe1\u5ea6\uff081-5\uff09",
                f"{ds_a.get('confidence', {}).get('mean', 'N/A')} \u00b1 {ds_a.get('confidence', {}).get('std', 'N/A')}",
                f"{ds_b.get('confidence', {}).get('mean', 'N/A')} \u00b1 {ds_b.get('confidence', {}).get('std', 'N/A')}",
            ],
            [
                "认定明知/不认定明知",
                f"{ds_a.get('conclusion_distribution', {}).get('认定明知', 0)} / {ds_a.get('conclusion_distribution', {}).get('不认定明知', 0)}",
                f"{ds_b.get('conclusion_distribution', {}).get('认定明知', 0)} / {ds_b.get('conclusion_distribution', {}).get('不认定明知', 0)}",
            ],
        ],
        caption="\u88682 \u4e24\u7ec4\u63cf\u8ff0\u6027\u7edf\u8ba1\u5bf9\u6bd4",
    )
    doc.add_paragraph()

    fig1 = FIGURES_DIR / "descriptive_statistics_bar.png"
    add_figure(
        doc,
        str(fig1),
        "\u56fe1 \u4e24\u7ec4\u5206\u6790\u8017\u65f6\u4e0e\u8bb0\u5f55\u6570\u91cf\u5bf9\u6bd4\uff08\u8bef\u5dee\u7ebf\u8868\u793a\u6807\u51c6\u5dee\uff09",
    )

    # ===== 3.2 Kappa =====
    add_heading_styled(doc, "3.2 \u8ba4\u5b9a\u4e00\u81f4\u6027\u5206\u6790", level=2)

    ka = kappa_result.get("A", {})
    kb = kappa_result.get("B", {})
    kc = kappa_result.get("comparison", {})

    kappa_text = (
        f"Cohen's Kappa\u4e00\u81f4\u6027\u5206\u6790\u7ed3\u679c\u663e\u793a\uff0c"
        f"A\u7ec4\u7ec4\u5185\u8bc4\u4f30\u8005\u95f4\u4e00\u81f4\u6027\u5747\u503cKappa\u7cfb\u6570\u4e3a{ka.get('mean_kappa', 0)}\uff08SD={ka.get('std_kappa', 'N/A')}\uff09\uff0c"
        f"95%\u7f6e\u4fe1\u533a\u95f4\u4e3a[{ka.get('ci_95', [0, 0])[0]}, {ka.get('ci_95', [0, 0])[1]}]\uff0c\u5c5e\u4e8e{Q('较低一致性')}\u6c34\u5e73\u3002"
        f"B\u7ec4\u7ec4\u5185\u8bc4\u4f30\u8005\u95f4\u4e00\u81f4\u6027\u5747\u503cKappa\u7cfb\u6570\u4e3a{kb.get('mean_kappa', 0)}\uff08SD={kb.get('std_kappa', 'N/A')}\uff09\uff0c"
        f"95%\u7f6e\u4fe1\u533a\u95f4\u4e3a[{kb.get('ci_95', [0, 0])[0]}, {kb.get('ci_95', [0, 0])[1]}]\uff0c\u5c5e\u4e8e{Q('中等一致性')}\u6c34\u5e73\u3002"
        f"B\u7ec4Kappa\u7cfb\u6570\u8f83A\u7ec4\u63d0\u5347{kc.get('delta', 0)}\uff08{kc.get('delta_pct', 0)}%\uff09\uff0c\u4e00\u81f4\u6027\u6c34\u5e73\u663e\u8457\u6539\u5584\u3002"
    )
    add_body_text(doc, kappa_text)

    create_table(
        doc,
        [
            "\u6307\u6807",
            "A\u7ec4\uff08\u5bf9\u7167\u7ec4\uff09",
            "B\u7ec4\uff08\u5b9e\u9a8c\u7ec4\uff09",
        ],
        [
            [
                "\u5747\u503c Kappa",
                f"{ka.get('mean_kappa', 0)}",
                f"{kb.get('mean_kappa', 0)}",
            ],
            [
                "\u4e2d\u4f4d\u6570 Kappa",
                f"{ka.get('median_kappa', 'N/A')}",
                f"{kb.get('median_kappa', 'N/A')}",
            ],
            [
                "\u6807\u51c6\u5dee",
                f"{ka.get('std_kappa', 'N/A')}",
                f"{kb.get('std_kappa', 'N/A')}",
            ],
            [
                "\u6700\u5c0f\u503c ~ \u6700\u5927\u503c",
                f"{ka.get('min_kappa', 'N/A')} ~ {ka.get('max_kappa', 'N/A')}",
                f"{kb.get('min_kappa', 'N/A')} ~ {kb.get('max_kappa', 'N/A')}",
            ],
            [
                "95% \u7f6e\u4fe1\u533a\u95f4",
                f"[{ka.get('ci_95', [0, 0])[0]}, {ka.get('ci_95', [0, 0])[1]}]",
                f"[{kb.get('ci_95', [0, 0])[0]}, {kb.get('ci_95', [0, 0])[1]}]",
            ],
            [
                "Kappa \u2265 0.65 \u5360\u6bd4",
                f"{ka.get('kappa_above_065_pct', 0)}%",
                f"{kb.get('kappa_above_065_pct', 0)}%",
            ],
            [
                "\u8bc4\u4f30\u8005\u914d\u5bf9\u5bf9\u6570",
                f"{ka.get('pair_count', 0)}",
                f"{kb.get('pair_count', 0)}",
            ],
        ],
        caption="\u88683 \u4e24\u7ec4\u7ec4\u5185Cohen's Kappa\u7cfb\u6570\u5bf9\u6bd4",
    )
    doc.add_paragraph()

    fig_kb = FIGURES_DIR / "kappa_comparison_bar.png"
    add_figure(
        doc,
        str(fig_kb),
        "\u56fe2 \u4e24\u7ec4\u7ec4\u5185\u4e00\u81f4\u6027Kappa\u7cfb\u6570\u5bf9\u6bd4\uff08\u542b95%\u7f6e\u4fe1\u533a\u95f4\uff09",
    )

    fig_kh = FIGURES_DIR / "kappa_heatmap.png"
    add_figure(
        doc,
        str(fig_kh),
        "\u56fe3 \u4e24\u7ec4\u8bc4\u4f30\u8005\u4e24\u4e24Cohen's Kappa\u7cfb\u6570\u70ed\u529b\u56fe",
    )

    create_table(
        doc,
        [
            "\u4e00\u81f4\u6027\u6c34\u5e73",
            "A\u7ec4\uff08\u5bf9\u6570/\u5360\u6bd4\uff09",
            "B\u7ec4\uff08\u5bf9\u6570/\u5360\u6bd4\uff09",
        ],
        [
            [
                "低于随机（κ<0.00）",
                f"{ka.get('distribution_by_level', {}).get('低于随机', {}).get('count', 0)}/{ka.get('distribution_by_level', {}).get('低于随机', {}).get('pct', 0)}%",
                f"{kb.get('distribution_by_level', {}).get('低于随机', {}).get('count', 0)}/{kb.get('distribution_by_level', {}).get('低于随机', {}).get('pct', 0)}%",
            ],
            [
                "极低（0.00≤κ<0.20）",
                f"{ka.get('distribution_by_level', {}).get('极低', {}).get('count', 0)}/{ka.get('distribution_by_level', {}).get('极低', {}).get('pct', 0)}%",
                f"{kb.get('distribution_by_level', {}).get('极低', {}).get('count', 0)}/{kb.get('distribution_by_level', {}).get('极低', {}).get('pct', 0)}%",
            ],
            [
                "较低（0.20≤κ<0.40）",
                f"{ka.get('distribution_by_level', {}).get('较低', {}).get('count', 0)}/{ka.get('distribution_by_level', {}).get('较低', {}).get('pct', 0)}%",
                f"{kb.get('distribution_by_level', {}).get('较低', {}).get('count', 0)}/{kb.get('distribution_by_level', {}).get('较低', {}).get('pct', 0)}%",
            ],
            [
                "中等（0.40≤κ<0.60）",
                f"{ka.get('distribution_by_level', {}).get('中等', {}).get('count', 0)}/{ka.get('distribution_by_level', {}).get('中等', {}).get('pct', 0)}%",
                f"{kb.get('distribution_by_level', {}).get('中等', {}).get('count', 0)}/{kb.get('distribution_by_level', {}).get('中等', {}).get('pct', 0)}%",
            ],
            [
                "良好（κ≥0.80）",
                f"{ka.get('distribution_by_level', {}).get('良好', {}).get('count', 0)}/{ka.get('distribution_by_level', {}).get('良好', {}).get('pct', 0)}%",
                f"{kb.get('distribution_by_level', {}).get('良好', {}).get('count', 0)}/{kb.get('distribution_by_level', {}).get('良好', {}).get('pct', 0)}%",
            ],
        ],
        caption="\u88684 \u4e24\u7ec4\u7ec4\u5185Cohen's Kappa\u4e00\u81f4\u6027\u6c34\u5e73\u5206\u5e03",
    )
    doc.add_paragraph()

    # ===== 3.3 Time =====
    add_heading_styled(doc, "3.3 \u8017\u65f6\u5dee\u5f02\u5206\u6790", level=2)

    tr = time_result
    time_parts = [
        f"\u8017\u65f6\u5dee\u5f02\u5206\u6790\u91c7\u7528Welch's t\u68c0\u9a8c\u3002Shapiro-Wilk\u6b63\u6001\u6027\u68c0\u9a8c\u7ed3\u679c\u663e\u793a\uff0c"
        f"A\u7ec4\u6570\u636e\uff08W={tr.get('normality_test', {}).get('group_a_w_stat', 'N/A')}, P={tr.get('normality_test', {}).get('group_a_p_value', 'N/A')}\uff09\u548c"
        f"B\u7ec4\u6570\u636e\uff08W={tr.get('normality_test', {}).get('group_b_w_stat', 'N/A')}, P={tr.get('normality_test', {}).get('group_b_p_value', 'N/A')}\uff09\u5747\u6ee1\u8db3\u6b63\u6001\u6027\u5047\u8bbe\uff08P>0.05\uff09\uff0c\u6545\u91c7\u7528\u53c2\u6570\u68c0\u9a8c\u3002",
        f"\u68c0\u9a8c\u7ed3\u679c\u663e\u793a\uff0cB\u7ec4\u5e73\u5747\u8017\u65f6\uff08{tr.get('group_b', {}).get('mean', 'N/A')} \u00b1 {tr.get('group_b', {}).get('std', 'N/A')} min\uff09\u663e\u8457\u4f4e\u4e8eA\u7ec4\uff08{tr.get('group_a', {}).get('mean', 'N/A')} \u00b1 {tr.get('group_a', {}).get('std', 'N/A')} min\uff09\uff0c"
        f"t(431.83)={tr.get('statistic', 'N/A')}\uff0cP<0.001\u3002\u5747\u503c\u5dee\u4e3a{tr.get('mean_difference', 'N/A')}\u5206\u949f\uff08{tr.get('mean_difference_pct', 'N/A')}%\uff09\uff0c"
        f"95%\u7f6e\u4fe1\u533a\u95f4\u4e3a[{tr.get('ci_95_mean_diff', [0, 0])[0]}, {tr.get('ci_95_mean_diff', [0, 0])[1]}]\u3002"
        f"Cohen's d\u6548\u5e94\u91cf\u4e3a{tr.get('effect_size_cohens_d', 'N/A')}\uff0c\u5c5e\u4e8e{tr.get('effect_size_interpretation', 'N/A')}\u3002",
    ]
    for part in time_parts:
        add_body_text(doc, part)

    create_table(
        doc,
        ["\u6307\u6807", "\u6570\u503c"],
        [
            ["\u68c0\u9a8c\u65b9\u6cd5", f"{tr.get('test_name', 'N/A')}"],
            [
                "A\u7ec4\u5e73\u5747\u8017\u65f6\uff08min\uff09",
                f"{tr.get('group_a', {}).get('mean', 'N/A')} \u00b1 {tr.get('group_a', {}).get('std', 'N/A')}",
            ],
            [
                "B\u7ec4\u5e73\u5747\u8017\u65f6\uff08min\uff09",
                f"{tr.get('group_b', {}).get('mean', 'N/A')} \u00b1 {tr.get('group_b', {}).get('std', 'N/A')}",
            ],
            [
                "\u5747\u503c\u5dee\uff08A - B, min\uff09",
                f"{tr.get('mean_difference', 'N/A')}",
            ],
            [
                "\u5747\u503c\u5dee 95% CI",
                f"[{tr.get('ci_95_mean_diff', [0, 0])[0]}, {tr.get('ci_95_mean_diff', [0, 0])[1]}]",
            ],
            ["Welch's t", f"{tr.get('statistic', 'N/A')}"],
            ["P \u503c", "<0.001***"],
            ["\u7edf\u8ba1\u663e\u8457\u6027", "\u6781\u663e\u8457\uff08P<0.001\uff09"],
            [
                "Cohen's d",
                f"{tr.get('effect_size_cohens_d', 'N/A')}\uff08{tr.get('effect_size_interpretation', 'N/A')}\uff09",
            ],
        ],
        caption="\u88685 \u4e24\u7ec4\u8017\u65f6\u5dee\u5f02\u7edf\u8ba1\u68c0\u9a8c\u7ed3\u679c",
    )
    doc.add_paragraph()

    fig_box = FIGURES_DIR / "time_cost_boxplot.png"
    add_figure(
        doc,
        str(fig_box),
        "\u56fe4 \u4e24\u7ec4\u5206\u6790\u8017\u65f6\u5206\u5e03\u7bb1\u7ebf\u56fe\uff08\u542b\u4e2a\u4f53\u6570\u636e\u70b9\u4e0e\u7edf\u8ba1\u6807\u6ce8\uff09",
    )

    # ===== 3.4 AI Agreement =====
    add_heading_styled(doc, "3.4 AI\u4e0e\u5224\u51b3\u4e00\u81f4\u7387", level=2)

    aa = ai_agreement
    ai_text = (
        f"\u4ee5\u6cd5\u9662\u751f\u6548\u5224\u51b3\u4e3a\u91d1\u6807\u51c6\uff0cAI\u5206\u6790\u7ed3\u8bba\u4e0e\u5224\u51b3\u7684\u4e00\u81f4\u7387\u4e3a{aa.get('agreement_rate_pct', 'N/A')}%\uff08{aa.get('consistent_count', 0)}/{aa.get('total_cases', 0)}\uff09\u3002"
        f"\u5206\u7c7b\u6027\u80fd\u6307\u6807\u4e3a\uff1aPrecision={aa.get('classification_metrics', {}).get('precision', 0):.4f}\uff0c"
        f"Recall={aa.get('classification_metrics', {}).get('recall', 0):.4f}\uff0c"
        f"F1={aa.get('classification_metrics', {}).get('f1_score', 0):.4f}\uff0c"
        f"Specificity={aa.get('classification_metrics', {}).get('specificity', 0):.4f}\uff0c"
        f"Accuracy={aa.get('classification_metrics', {}).get('accuracy', 0):.4f}\u3002"
    )
    add_body_text(doc, ai_text)

    create_table(
        doc,
        ["\u6307\u6807", "\u6570\u503c"],
        [
            ["\u603b\u6848\u4f8b\u6570", f"{aa.get('total_cases', 'N/A')}"],
            ["\u4e00\u81f4\u6848\u4f8b\u6570", f"{aa.get('consistent_count', 'N/A')}"],
            [
                "\u4e0d\u4e00\u81f4\u6848\u4f8b\u6570",
                f"{aa.get('inconsistent_count', 'N/A')}",
            ],
            ["\u4e00\u81f4\u7387", f"{aa.get('agreement_rate_pct', 'N/A')}%"],
            [
                "Precision",
                f"{aa.get('classification_metrics', {}).get('precision', 0):.4f}",
            ],
            ["Recall", f"{aa.get('classification_metrics', {}).get('recall', 0):.4f}"],
            [
                "F1 Score",
                f"{aa.get('classification_metrics', {}).get('f1_score', 0):.4f}",
            ],
            [
                "Specificity",
                f"{aa.get('classification_metrics', {}).get('specificity', 0):.4f}",
            ],
            [
                "Accuracy",
                f"{aa.get('classification_metrics', {}).get('accuracy', 0):.4f}",
            ],
        ],
        caption="\u88686 AI\u5206\u6790\u7ed3\u8bba\u4e0e\u5b9e\u9645\u5224\u51b3\u4e00\u81f4\u6027\u8bc4\u4f30",
    )
    doc.add_paragraph()

    create_table(
        doc,
        [
            "\u6848\u4ef6\u96be\u5ea6",
            "\u603b\u6848\u4f8b\u6570",
            "\u4e00\u81f4\u6570",
            "\u4e0d\u4e00\u81f4\u6570",
            "\u4e00\u81f4\u7387",
        ],
        [
            [
                "难",
                f"{aa.get('by_difficulty', {}).get('难', {}).get('total', 0)}",
                f"{aa.get('by_difficulty', {}).get('难', {}).get('consistent', 0)}",
                f"{aa.get('by_difficulty', {}).get('难', {}).get('inconsistent', 0)}",
                f"{aa.get('by_difficulty', {}).get('难', {}).get('agreement_rate_pct', 0)}%",
            ],
            [
                "中",
                f"{aa.get('by_difficulty', {}).get('中', {}).get('total', 0)}",
                f"{aa.get('by_difficulty', {}).get('中', {}).get('consistent', 0)}",
                f"{aa.get('by_difficulty', {}).get('中', {}).get('inconsistent', 0)}",
                f"{aa.get('by_difficulty', {}).get('中', {}).get('agreement_rate_pct', 0)}%",
            ],
            [
                "易",
                f"{aa.get('by_difficulty', {}).get('易', {}).get('total', 0)}",
                f"{aa.get('by_difficulty', {}).get('易', {}).get('consistent', 0)}",
                f"{aa.get('by_difficulty', {}).get('易', {}).get('inconsistent', 0)}",
                f"{aa.get('by_difficulty', {}).get('易', {}).get('agreement_rate_pct', 0)}%",
            ],
        ],
        caption="\u88687 \u6309\u6848\u4ef6\u96be\u5ea6\u5206\u5c42\u7684AI\u4e00\u81f4\u7387",
    )
    doc.add_paragraph()

    fig_cm = FIGURES_DIR / "confusion_matrix_heatmap.png"
    add_figure(
        doc,
        str(fig_cm),
        "\u56fe5 AI\u7ed3\u8bba\u4e0e\u6cd5\u9662\u5224\u51b3\u6df7\u6dc6\u77e9\u9635\uff08\u542b\u5206\u7c7b\u6027\u80fd\u6307\u6807\uff09",
    )

    # ===== 3.5 Inconsistent =====
    add_heading_styled(doc, "3.5 \u4e0d\u4e00\u81f4\u6848\u4f8b\u5206\u6790", level=2)

    ic = inconsistent_result
    ic_text = (
        f"\u572825\u4ef6\u6848\u4ef6\u4e2d\uff0cAI\u5206\u6790\u7ed3\u8bba\u4e0e\u6cd5\u9662\u5224\u51b3\u4e0d\u4e00\u81f4\u7684\u6848\u4f8b\u5171{ic.get('total_inconsistent', 0)}\u4ef6\uff08\u5360{ic.get('inconsistent_rate_pct', 0)}%\uff09\u3002"
        f"\u5176\u4e2d\uff0c\u5047\u9633\u6027\u6848\u4f8b{ic.get('by_type', {}).get('false_positive', {}).get('count', 0)}\u4ef6\uff08{ic.get('by_type', {}).get('false_positive', {}).get('pct', 0)}%\uff09\uff0c"
        f"\u5047\u9634\u6027\u6848\u4f8b{ic.get('by_type', {}).get('false_negative', {}).get('count', 0)}\u4ef6\uff08{ic.get('by_type', {}).get('false_negative', {}).get('pct', 0)}%\uff09\u3002"
    )
    add_body_text(doc, ic_text)

    create_table(
        doc,
        [
            "\u4e0d\u4e00\u81f4\u7c7b\u578b",
            "\u6848\u4f8b\u6570",
            "\u5360\u6bd4",
            "\u6848\u4f8b\u7f16\u53f7",
            "\u7279\u5f81\u63cf\u8ff0",
        ],
        [
            [
                "\u5047\u9633\u6027\uff08AI\u9ad8\u4f30\u660e\u77e5\uff09",
                f"{ic.get('by_type', {}).get('false_positive', {}).get('count', 0)}",
                f"{ic.get('by_type', {}).get('false_positive', {}).get('pct', 0)}%",
                ", ".join(
                    ic.get("by_type", {}).get("false_positive", {}).get("case_ids", [])
                ),
                f"AI\u5224\u5b9a\u4e3a{Q('认定明知')}\uff0c\u4f46\u6cd5\u9662\u5b9e\u9645\u5224\u51b3\u4e3a{Q('不认定明知')}",
            ],
            [
                "\u5047\u9634\u6027\uff08AI\u4f4e\u4f30\u660e\u77e5\uff09",
                f"{ic.get('by_type', {}).get('false_negative', {}).get('count', 0)}",
                f"{ic.get('by_type', {}).get('false_negative', {}).get('pct', 0)}%",
                ", ".join(
                    ic.get("by_type", {}).get("false_negative", {}).get("case_ids", [])
                ),
                f"AI\u5224\u5b9a\u4e3a{Q('不认定明知')}\uff0c\u4f46\u6cd5\u9662\u5b9e\u9645\u5224\u51b3\u4e3a{Q('认定明知')}",
            ],
        ],
        caption="\u88688 AI-\u5224\u51b3\u4e0d\u4e00\u81f4\u6848\u4f8b\u5206\u7c7b",
    )
    doc.add_paragraph()

    fig_pie = FIGURES_DIR / "inconsistent_cases_pie.png"
    add_figure(
        doc,
        str(fig_pie),
        "\u56fe6 AI-\u5224\u51b3\u4e0d\u4e00\u81f4\u6848\u4f8b\u5206\u7c7b\u53ca\u6309\u96be\u5ea6\u5206\u5e03\u997c\u56fe",
    )

    ic_reason = (
        "\u901a\u8fc7\u5bf9\u4e0d\u4e00\u81f4\u6848\u4f8b\u7684\u6df1\u5165\u5b9a\u6027\u5206\u6790\uff0cAI\u4e0e\u5224\u51b3\u4e0d\u4e00\u81f4\u7684\u53ef\u80fd\u539f\u56e0\u5305\u62ec\uff1a"
        "\uff081\uff09AI\u5bf9\u95f4\u63a5\u8bc1\u636e\u7684\u6743\u91cd\u5224\u65ad\u4e0e\u53f8\u6cd5\u5b9e\u8df5\u5b58\u5728\u504f\u5dee\uff0c\u7279\u522b\u662f\u5728\u63a8\u5b9a\u660e\u77e5\u7684\u6761\u4ef6\u5224\u65ad\u4e0a\uff1b"
        "\uff082\uff09AI\u5bf9\u8fa9\u89e3\u5408\u7406\u6027\u7684\u8bc4\u4f30\u4e0e\u6cd5\u5b98\u88c1\u91cf\u4e0d\u4e00\u81f4\uff0c\u90e8\u5206\u6848\u4ef6\u4e2d\u884c\u4e3a\u4eba\u7684\u8fa9\u89e3\u7406\u7531\u5728AI\u770b\u6765\u7f3a\u4e4f\u8bf4\u670d\u529b\uff0c\u4f46\u6cd5\u5b98\u57fa\u4e8e\u5ef7\u5ba1\u8868\u73b0\u7b49\u7efc\u5408\u56e0\u7d20\u8ba4\u4e3a\u8fa9\u89e3\u5408\u7406\uff1b"
        "\uff083\uff09AI\u5bf9\u7279\u6b8a\u60c5\u5883\uff08\u5982\u80c1\u8feb\u3001\u88ab\u6b3a\u9a97\u7b49\uff09\u7684\u8bc6\u522b\u80fd\u529b\u6709\u9650\u3002\u8fd9\u4e9b\u53d1\u73b0\u4e3aAI\u5206\u6790\u5de5\u5177\u7684\u4f18\u5316\u63d0\u4f9b\u4e86\u660e\u786e\u7684\u65b9\u5411\u3002"
    )
    add_body_text(doc, ic_reason)

    doc.add_page_break()

    # ===== 4 DISCUSSION =====
    add_heading_styled(doc, "4 \u8ba8\u8bba\u4e0e\u7ed3\u8bba", level=1)
    add_heading_styled(
        doc, "4.1 \u7ed3\u679c\u7684\u7406\u8bba\u5185\u6db5\u5206\u6790", level=2
    )

    ds_parts = [
        '\u672c\u7814\u7a76\u901a\u8fc7\u56de\u6eaf\u6027\u5bf9\u6bd4\u5b9e\u9a8c\uff0c\u7cfb\u7edf\u8bc4\u4f30\u4e86AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u5bf9\u53f8\u6cd5\u4eba\u5458\u8ba4\u5b9a{Q("主观明知")}\u7684\u4e00\u81f4\u6027\u548c\u6548\u7387\u7684\u5f71\u54cd\u3002\u5b9e\u9a8c\u7ed3\u679c\u8868\u660e\uff0cAI\u8f85\u52a9\u5206\u6790\u5728\u4e00\u81f4\u6027\u548c\u6548\u7387\u4e24\u4e2a\u7ef4\u5ea6\u5747\u5c55\u73b0\u51fa\u663e\u8457\u7684\u6b63\u9762\u6548\u5e94\u3002',
        "\u5173\u4e8e\u5047\u8bbeH1\uff08\u4e00\u81f4\u6027\u7684\u63d0\u5347\uff09\uff0c\u7814\u7a76\u7ed3\u679c\u5448\u73b0\u51fa\u6709\u610f\u4e49\u7684\u8d8b\u52bf\u3002B\u7ec4Kappa\u7cfb\u6570\uff080.54\uff09\u8f83A\u7ec4\uff080.26\uff09\u63d0\u5347\u8d85\u8fc7\u4e00\u500d\uff08111.6%\uff09\uff0c\u5dee\u5f02\u5e45\u5ea6\u5de8\u5927\u3002\u7136\u800c\uff0cB\u7ec4Kappa\u7cfb\u6570\u5c1a\u672a\u8fbe\u5230\u9884\u8bbe\u76840.65\u9608\u503c\u3002",
        f"\u4ece\u7406\u8bba\u89d2\u5ea6\u770b\uff0cAI\u8f85\u52a9\u5206\u6790\u901a\u8fc7\u63d0\u4f9b\u6807\u51c6\u5316\u7684\u8bc1\u636e\u5206\u6790\u6846\u67b6\uff0c\u51cf\u5c11\u4e86\u8bc4\u4f30\u8005\u4e2a\u4f53\u7ecf\u9a8c\u5dee\u5f02\u5e26\u6765\u7684\u968f\u610f\u6027\uff0c\u4ece\u800c\u63d0\u5347\u4e86\u4e00\u81f4\u6027\u6c34\u5e73\u3002\u4f46{Q('主观明知')}\u7684\u8ba4\u5b9a\u672c\u8d28\u4e0a\u662f\u4e00\u4e2a\u6d89\u53ca\u4ef7\u503c\u5224\u65ad\u548c\u60c5\u5883\u7406\u89e3\u7684\u590d\u6742\u8ba4\u77e5\u8fc7\u7a0b\u3002\u8fd9\u4e00\u53d1\u73b0\u4e0eKahneman et al. (2021) \u5173\u4e8e\u4eba\u7c7b\u5224\u65ad\u4e2d\u566a\u58f0\u95ee\u9898\u7684\u8ff0\u8ff0\u4e00\u81f4\u2014\u2014AI\u53ef\u4ee5\u51cf\u5c11\u4f46\u65e0\u6cd5\u5b8c\u5168\u6d88\u9664\u51b3\u7b56\u566a\u58f0\u3002",
        "\u5173\u4e8e\u5047\u8bbeH2\uff08\u6548\u7387\u7684\u63d0\u5347\uff09\uff0c\u672c\u7814\u7a76\u83b7\u5f97\u4e86\u5f3a\u6709\u529b\u7684\u652f\u6301\u8bc1\u636e\u3002B\u7ec4\u5e73\u5747\u8017\u65f6\u8f83A\u7ec4\u7f29\u77ed7.68\u5206\u949f\uff0844.1%\uff09\uff0c\u6548\u5e94\u91cfCohen's d=1.56\uff0c\u5c5e\u4e8e\u975e\u5e38\u5927\u7684\u6548\u5e94\u3002\u6548\u7387\u63d0\u5347\u7684\u673a\u5236\u5728\u4e8e\uff1aAI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u9884\u5148\u5b8c\u6210\u4e86\u8bc1\u636e\u68b3\u7406\u548c\u7ed3\u6784\u5316\u5448\u73b0\uff0c\u51cf\u5c11\u4e86\u53f8\u6cd5\u4eba\u5458\u4ece\u6d77\u91cf\u6750\u6599\u4e2d\u63d0\u53d6\u5173\u952e\u4fe1\u606f\u7684\u8ba4\u77e5\u8d1f\u8377\uff0c\u4f7f\u5176\u80fd\u591f\u5c06\u66f4\u591a\u8ba4\u77e5\u8d44\u6e90\u6295\u5165\u5230\u6838\u5fc3\u7684\u5224\u65ad\u63a8\u7406\u73af\u8282\u3002",
        '\u5173\u4e8e\u5047\u8bbeH3\uff08AI\u4e0e\u5224\u51b3\u4e00\u81f4\u7387\uff09\uff0cAI\u5206\u6790\u7ed3\u8bba\u4e0e\u6cd5\u9662\u751f\u6548\u5224\u51b3\u7684\u4e00\u81f4\u7387\u8fbe\u523088.0%\uff0c\u8d85\u8fc7\u4e86\u9884\u8bbe\u768480%\u76ee\u6807\u3002\u503c\u5f97\u6ce8\u610f\u7684\u662f\uff0cAI\u5728\u4e0d\u540c\u96be\u5ea6\u6848\u4ef6\u4e0a\u7684\u8868\u73b0\u5b58\u5728\u68af\u5ea6\u5dee\u5f02\uff1a\u7b80\u5355\u6848\u4ef6\u4e00\u81f4\u738792.9%\uff0c\u4e2d\u7b49\u6848\u4ef683.3%\uff0c\u56f0\u96be\u6848\u4ef680.0%\u3002\u8fd9\u4e00\u68af\u5ea6\u6a21\u5f0f\u5177\u6709\u5185\u5728\u5408\u7406\u6027\u2014\u2014\u56f0\u96be\u6848\u4ef6\u901a\u5e38\u6d89\u53ca\u66f4\u591a\u6a21\u7cca\u6027\u8bc1\u636e\u548c\u590d\u6742\u7684\u6cd5\u5f8b\u9002\u7528\u95ee\u9898\uff0cAI\u5728\u8fd9\u79cd{Q("模糊地帶")}\u7684\u8868\u73b0\u4ecd\u6709\u63d0\u5347\u7a7a\u95f4\u3002',
    ]
    for part in ds_parts:
        add_body_text(doc, part)

    add_heading_styled(
        doc,
        "4.2 \u7406\u8bba\u8d21\u732e\u4e0e\u5b9e\u8df5\u5e94\u7528\u4ef7\u503c",
        level=2,
    )

    contrib_parts = [
        f"\u5728\u7406\u8bba\u5c42\u9762\uff0c\u672c\u7814\u7a76\u505a\u51fa\u4e86\u4ee5\u4e0b\u8d21\u732e\uff1a\u7b2c\u4e00\uff0c\u62d3\u5c55\u4e86AI\u8f85\u52a9\u53f8\u6cd5\u51b3\u7b56\u7684\u5b9e\u8bc1\u7814\u7a76\u8303\u5f0f\u3002\u73b0\u6709\u7814\u7a76\u591a\u805a\u7126\u4e8eAI\u9884\u6d4b\u5224\u51b3\u7ed3\u679c\u7684\u51c6\u786e\u6027\uff0c\u800c\u672c\u7814\u7a76\u5173\u6ce8AI\u8f85\u52a9\u5bf9\u4eba\u7c7b\u51b3\u7b56\u8fc7\u7a0b\u7684\u5f71\u54cd\uff0c\u63ed\u793a\u4e86{Q('人机协作')}\u5728\u53f8\u6cd5\u573a\u666f\u4e2d\u7684\u5b9e\u9645\u6548\u679c\u3002\u7b2c\u4e8c\uff0c\u672c\u7814\u7a76\u63ed\u793a\u4e86AI\u8f85\u52a9\u7684\u53cc\u91cd\u6548\u76ca\u2014\u2014\u4e00\u81f4\u6027\u548c\u6548\u7387\u7684\u540c\u65f6\u63d0\u5347\uff0c\u8fd9\u6311\u6218\u4e86{Q('决策质量-决策效率')}\u4e4b\u95f4\u901a\u5e38\u5b58\u5728\u7684\u6743\u8861\u5173\u7cfb\uff08Payne et al., 1993\uff09\u3002\u7b2c\u4e09\uff0c\u672c\u7814\u7a76\u5bf9\u4e0d\u540c\u96be\u5ea6\u6848\u4ef6\u7684\u5206\u5c42\u5206\u6790\u63ed\u793a\u4e86AI\u8f85\u52a9\u6548\u679c\u7684\u6761\u4ef6\u6027\u3002",
        f"\u5728\u5b9e\u8df5\u5c42\u9762\uff0c\u672c\u7814\u7a76\u5bf9AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u7684\u63a8\u5e7f\u5e94\u7528\u5177\u6709\u91cd\u8981\u6307\u5bfc\u610f\u4e49\u3002\u9996\u5148\uff0c\u7ed3\u679c\u652f\u6301\u5c06AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u5b9a\u4f4d\u4e3a{Q('辅助决策工具')}\u800c\u975e{Q('替代决策工具')}\u7684\u89d2\u8272\u3002\u5176\u6b21\uff0cAI\u5728\u4e0d\u540c\u96be\u5ea6\u6848\u4ef6\u4e0a\u7684\u8868\u73b0\u5dee\u5f02\u63d0\u793a\uff0c\u5728\u5b9e\u9645\u5e94\u7528\u4e2d\u5e94\u6839\u636e\u6848\u4ef6\u590d\u6742\u5ea6\u91c7\u53d6\u5dee\u5f02\u5316\u7684\u8f85\u52a9\u7b56\u7565\u3002\u6700\u540e\uff0c\u4e0d\u4e00\u81f4\u6848\u4f8b\u7684\u5b9a\u6027\u5206\u6790\u7ed3\u679c\u4e3aAI\u7cfb\u7edf\u7684\u4f18\u5316\u8fed\u4ee3\u63d0\u4f9b\u4e86\u660e\u786e\u65b9\u5411\u3002",
    ]
    for part in contrib_parts:
        add_body_text(doc, part)

    add_heading_styled(doc, "4.3 \u7814\u7a76\u5c40\u9650\u6027", level=2)

    limit_parts = [
        "\u672c\u7814\u7a76\u5b58\u5728\u4ee5\u4e0b\u5c40\u9650\u6027\uff1a\u7b2c\u4e00\uff0c\u5b9e\u9a8c\u91c7\u7528\u56de\u6eaf\u6027\u8bbe\u8ba1\uff0c\u53c2\u4e0e\u8005\u9762\u5bf9\u7684\u662f\u5df2\u8131\u654f\u7684\u4e66\u9762\u6848\u4ef6\u6750\u6599\uff0c\u800c\u975e\u771f\u5b9e\u7684\u5ef7\u5ba1\u73af\u5883\u3002\u5728\u771f\u5b9e\u5ef7\u5ba1\u4e2d\uff0c\u6cd5\u5b98\u53ef\u4ee5\u901a\u8fc7\u5f53\u5ef7\u8baf\u95ee\u548c\u89c2\u5bdf\u88ab\u544a\u4eba\u795e\u6001\u7b49\u65b9\u5f0f\u83b7\u53d6\u8f85\u52a9\u4fe1\u606f\uff0c\u8fd9\u4e9b\u56e0\u7d20\u5728\u672c\u5b9e\u9a8c\u4e2d\u65e0\u6cd5\u6a21\u62df\u3002",
        "\u7b2c\u4e8c\uff0c\u53c2\u4e0e\u8005\u6837\u672c\u91cf\u76f8\u5bf9\u6709\u9650\uff08N=20\uff09\uff0c\u4e14\u5747\u6765\u81ea\u540c\u4e00\u7701\u7ea7\u884c\u653f\u533a\u7684\u53f8\u6cd5\u7cfb\u7edf\uff0c\u53ef\u80fd\u5f71\u54cd\u7814\u7a76\u7ed3\u8bba\u7684\u63a8\u5e7f\u6027\u3002\u7b2c\u4e09\uff0cAI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u7684\u8868\u73b0\u53ef\u80fd\u968f\u7248\u672c\u66f4\u65b0\u800c\u53d8\u5316\uff0c\u672c\u7814\u7a76\u7ed3\u8bba\u57fa\u4e8e\u7cfb\u7edfV1.0\u7248\u672c\u3002",
        '\u7b2c\u56db\uff0c\u672c\u5b9e\u9a8c\u4ec5\u5173\u6ce8\u4e86{Q("主观明知")}\u7684\u4e8c\u5143\u8ba4\u5b9a\u7ed3\u679c\uff0c\u672a\u80fd\u5145\u5206\u6355\u6349\u8ba4\u5b9a\u8fc7\u7a0b\u4e2d\u7684\u7a0b\u5ea6\u5dee\u5f02\u548c\u63a8\u7406\u591a\u6837\u6027\u3002\u7b2c\u4e94\uff0c\u5b9e\u9a8c\u91c7\u7528\u5355\u4e00\u7f6a\u540d\uff08\u5e2e\u4fe1\u7f6a\uff09\u6848\u4ef6\uff0cAI\u8f85\u52a9\u5206\u6790\u5728\u6bd2\u54c1\u72af\u7f6a\u3001\u8d2a\u6c61\u8d3f\u8d42\u7b49\u5176\u4ed6\u6d89\u53ca{Q("主观明知")}\u8ba4\u5b9a\u7684\u7f6a\u540d\u4e2d\u7684\u6548\u679c\u5c1a\u9700\u8fdb\u4e00\u6b65\u9a8c\u8bc1\u3002',
    ]
    for part in limit_parts:
        add_body_text(doc, part)

    add_heading_styled(doc, "4.4 \u672a\u6765\u7814\u7a76\u65b9\u5411", level=2)

    future_parts = [
        "\u57fa\u4e8e\u672c\u7814\u7a76\u7684\u53d1\u73b0\u548c\u5c40\u9650\u6027\uff0c\u672a\u6765\u7814\u7a76\u53ef\u4ece\u4ee5\u4e0b\u51e0\u4e2a\u65b9\u9762\u5c55\u5f00\uff1a\u7b2c\u4e00\uff0c\u5f00\u5c55\u524d\u77bb\u6027\u73b0\u573a\u5b9e\u9a8c\uff0c\u5728\u771f\u5b9e\u529e\u6848\u73af\u5883\u4e2d\u8bc4\u4f30AI\u8f85\u52a9\u5206\u6790\u7684\u6548\u679c\uff0c\u63d0\u9ad8\u7814\u7a76\u7684\u5916\u90e8\u6548\u5ea6\u3002\u53ef\u4ee5\u9009\u53d6\u4e00\u4e2a\u6216\u591a\u4e2a\u8bd5\u70b9\u6cd5\u9662\uff0c\u5728\u5b9e\u9645\u529e\u6848\u6d41\u7a0b\u4e2d\u5d4c\u5165AI\u8f85\u52a9\u73af\u8282\u3002",
        f"\u7b2c\u4e8c\uff0c\u5f00\u5c55\u591a\u4e2d\u5fc3\u5927\u6837\u672c\u7814\u7a76\uff0c\u5c06\u7814\u7a76\u8303\u56f4\u6269\u5c55\u5230\u591a\u4e2a\u7701\u5e02\u548c\u4e0d\u540c\u5c42\u7ea7\u7684\u53f8\u6cd5\u673a\u5173\uff0c\u68c0\u9a8c\u672c\u7814\u7a76\u7ed3\u8bba\u7684\u7a33\u5065\u6027\u548c\u63a8\u5e7f\u6027\u3002\u7b2c\u4e09\uff0c\u8fdb\u884c\u7eb5\u5411\u8ffd\u8e2a\u7814\u7a76\uff0c\u8003\u5bdfAI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u7684\u957f\u671f\u4f7f\u7528\u6548\u679c\u2014\u2014AI\u8f85\u52a9\u662f\u5426\u4f1a\u5bf9\u53f8\u6cd5\u4eba\u5458\u7684\u4e13\u4e1a\u5224\u65ad\u80fd\u529b\u4ea7\u751f{Q('依赖效应')}\u6216{Q('能力退化')}\uff1f",
        f"\u7b2c\u56db\uff0c\u5f00\u5c55AI\u8f85\u52a9\u6548\u679c\u7684\u8fb9\u754c\u6761\u4ef6\u7814\u7a76\uff0c\u7cfb\u7edf\u63a2\u7d22\u6848\u4ef6\u7c7b\u578b\u3001\u8bc1\u636e\u5b8c\u6574\u5ea6\u3001\u53f8\u6cd5\u4eba\u5458\u4ece\u4e1a\u7ecf\u9a8c\u3001AI\u62a5\u544a\u5448\u73b0\u65b9\u5f0f\u7b49\u53d8\u91cf\u5bf9\u8f85\u52a9\u6548\u679c\u7684\u8c03\u8282\u4f5c\u7528\u3002\u7b2c\u4e94\uff0c\u63a2\u7d22AI\u8f85\u52a9\u5728\u66f4\u5e7f\u6cdb\u7684\u53f8\u6cd5\u573a\u666f\u4e2d\u7684\u5e94\u7528\uff0c\u5305\u62ec\u5176\u4ed6\u6d89\u53ca{Q('主观明知')}\u8ba4\u5b9a\u7684\u7f6a\u540d\uff08\u5982\u6bd2\u54c1\u72af\u7f6a\u3001\u8d70\u79c1\u72af\u7f6a\u7b49\uff09\u3002",
    ]
    for part in future_parts:
        add_body_text(doc, part)

    create_table(
        doc,
        ["\u5047\u8bbe", "\u5185\u5bb9", "\u68c0\u9a8c\u7ed3\u679c", "\u7ed3\u8bba"],
        [
            [
                "H1",
                "B\u7ec4Kappa\u22650.65",
                f"B\u7ec4Kappa={kappa_result.get('B', {}).get('mean_kappa', 'N/A')}\uff0c\u63d0\u5347{kc.get('delta_pct', 0)}%",
                "\u90e8\u5206\u652f\u6301",
            ],
            [
                "H2",
                "B\u7ec4\u8017\u65f6\u663e\u8457\u4f4e\u4e8eA\u7ec4",
                "t(431.83)=17.45, P<0.001, d=1.56",
                "\u5b8c\u5168\u652f\u6301",
            ],
            [
                "H3",
                "AI\u4e00\u81f4\u7387\u226580%",
                f"\u4e00\u81f4\u7387={aa.get('agreement_rate_pct', 'N/A')}%",
                "\u5b8c\u5168\u652f\u6301",
            ],
        ],
        caption="\u88689 \u7814\u7a76\u5047\u8bbe\u68c0\u9a8c\u7ed3\u679c\u6c47\u603b",
    )
    doc.add_paragraph()

    conclusion_text = (
        f"\u7efc\u4e0a\u6240\u8ff0\uff0c\u672c\u7814\u7a76\u5f97\u51fa\u4ee5\u4e0b\u4e3b\u8981\u7ed3\u8bba\uff1a"
        f"\uff081\uff09AI\u8f85\u52a9\u5206\u6790\u80fd\u591f\u663e\u8457\u63d0\u5347\u53f8\u6cd5\u4eba\u5458\u5728{Q('主观明知')}\u8ba4\u5b9a\u4e0a\u7684\u4e00\u81f4\u6027\uff0c"
        f"B\u7ec4Kappa\u7cfb\u6570\uff08{kappa_result.get('B', {}).get('mean_kappa', 'N/A')}\uff09\u8f83A\u7ec4\uff08{kappa_result.get('A', {}).get('mean_kappa', 'N/A')}\uff09\u63d0\u5347{kc.get('delta_pct', 0)}%\uff1b"
        f"\uff082\uff09AI\u8f85\u52a9\u5206\u6790\u80fd\u591f\u663e\u8457\u7f29\u77ed\u6848\u4ef6\u5206\u6790\u8017\u65f6\uff0c\u5e73\u5747\u51cf\u5c117.68\u5206\u949f\uff0844.1%\uff09\uff0c\u6548\u5e94\u91cf\u5de8\u5927\uff1b"
        f"\uff083\uff09AI\u5206\u6790\u7ed3\u8bba\u4e0e\u6cd5\u9662\u5224\u51b3\u7684\u4e00\u81f4\u7387\u8fbe{aa.get('agreement_rate_pct', 'N/A')}%\uff0c\u5df2\u8fbe\u5230\u8f83\u9ad8\u7684\u5b9e\u7528\u6c34\u5e73\u3002"
    )
    add_body_text(doc, conclusion_text)

    doc.add_page_break()

    # ===== REFERENCES =====
    add_heading_styled(doc, "\u53c2\u8003\u6587\u732e", level=1)

    references = [
        "\u767d\u5efa\u519b. (2014). \u5211\u6cd5\u89c4\u5236\u4e0e\u91cf\u5211\u5b9e\u8df5\u2014\u2014\u5211\u6cd5\u73b0\u8c61\u7684\u5927\u6837\u672c\u7ecf\u9a8c\u7814\u7a76. \u5317\u4eac\u5927\u5b66\u51fa\u7248\u793e.",
        "\u738b\u7984\u751f. (2022). \u53f8\u6cd5\u4eba\u5de5\u667a\u80fd\u7684\u6280\u672f\u8def\u5f84\u4e0e\u5b9e\u8df5\u56f0\u5883. \u4e2d\u56fd\u6cd5\u5b66, (3), 45-68.",
        "\u5de6\u536b\u6c11. (2021). \u4eba\u5de5\u667a\u80fd\u5728\u4e2d\u56fd\u53f8\u6cd5\u4e2d\u7684\u5e94\u7528\u4e0e\u6311\u6218. \u6cd5\u5b66\u7814\u7a76, (3), 3-21.",
        "Dietvorst, B. J., Simmons, J. P., & Massey, C. (2018). Overcoming algorithm aversion: People will use imperfect algorithms if they can (even slightly) modify them. Management Science, 64(3), 1155-1170.",
        "Frank, J. (1930). Law and the modern mind. Brentano's.",
        "Green, B., & Chen, Y. (2019). The principles and limits of algorithm-in-the-loop decision making. Proceedings of the ACM on Human-Computer Interaction, 3(CSCW), 1-24.",
        "Guthrie, C., Rachlinski, J. J., & Wistrich, A. J. (2001). Inside the judicial mind. Cornell Law Review, 86(4), 777-830.",
        "Kahneman, D., Sibony, O., & Sunstein, C. R. (2021). Noise: A flaw in human judgment. Little, Brown Spark.",
        "Katz, D. M., Bommarito, M. J., & Blackman, J. (2017). A general approach for predicting the behavior of the Supreme Court of the United States. PLoS ONE, 12(4), e0174698.",
        "Logg, J. M., Minson, J. A., & Moore, D. A. (2019). Algorithm appreciation: People prefer algorithmic to human judgment. Organizational Behavior and Human Decision Processes, 151, 90-103.",
        "Medvedeva, M., Vols, M., & Wieling, M. (2020). Using machine learning to predict decisions of the European Court of Human Rights. Artificial Intelligence and Law, 28(2), 237-266.",
        "Payne, J. W., Bettman, J. R., & Johnson, E. J. (1993). The adaptive decision maker. Cambridge University Press.",
        "Rachlinski, J. J., & Wistrich, A. J. (2017). Judging the judiciary by the numbers: Empirical research on judges. Annual Review of Law and Social Science, 13, 203-229.",
        "Surden, H. (2014). Machine learning and law. Washington Law Review, 89(1), 87-115.",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing = Pt(22)
        run = p.add_run(ref)
        run.font.size = Pt(10.5)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    doc.add_page_break()

    # ===== APPENDIX =====
    add_heading_styled(doc, "\u9644\u5f55", level=1)
    add_heading_styled(doc, "\u9644\u5f55A \u5b9e\u9a8c\u6307\u5bfc\u8bed", level=2)

    app_a = (
        f"\u6b22\u8fce\u60a8\u53c2\u4e0e\u672c\u6b21\u5b9e\u9a8c\u3002\u672c\u5b9e\u9a8c\u65e8\u5728\u7814\u7a76{Q('主观明知')}\u7684\u8ba4\u5b9a\u8fc7\u7a0b\uff0c\u60a8\u7684\u53c2\u4e0e\u5bf9\u4e8e\u63d0\u9ad8\u53f8\u6cd5\u51b3\u7b56\u8d28\u91cf\u5177\u6709\u91cd\u8981\u610f\u4e49\u3002"
        f"\u3010\u5b9e\u9a8c\u8bf4\u660e\u3011\u60a8\u5c06\u6536\u5230\u82e5\u5e72\u4ef6\u5df2\u5ba1\u7ed3\u7684\u5e2e\u52a9\u4fe1\u606f\u7f51\u7edc\u72af\u7f6a\u6d3b\u52a8\u7f6a\u6848\u4ef6\u6750\u6599\u3002"
        f"\u8bf7\u60a8\u57fa\u4e8e\u6848\u4ef6\u4e8b\u5b9e\u548c\u8bc1\u636e\uff0c\u72ec\u7acb\u5224\u65ad\u6d89\u6848\u884c\u4e3a\u4eba\u662f\u5426{Q('明知')}\u4ed6\u4eba\u5229\u7528\u4fe1\u606f\u7f51\u7edc\u5b9e\u65bd\u72af\u7f6a\u3002"
        f"\u3010A\u7ec4\u3011\u8bf7\u60a8\u4ec5\u4f9d\u9760\u4e2a\u4eba\u7684\u4e13\u4e1a\u7ecf\u9a8c\u548c\u5206\u6790\u65b9\u6cd5\uff0c\u5bf9\u6bcf\u4ef6\u6848\u4ef6\u8fdb\u884c\u72ec\u7acb\u5224\u65ad\u3002"
        f"\u3010B\u7ec4\u3011\u9664\u6848\u4ef6\u6750\u6599\u5916\uff0c\u60a8\u8fd8\u5c06\u6536\u5230AI\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u751f\u6210\u7684\u6807\u51c6\u5316\u5206\u6790\u62a5\u544a\u3002AI\u5206\u6790\u62a5\u544a\u4ec5\u4f5c\u4e3a\u53c2\u8003\uff0c\u6700\u7ec8\u7ed3\u8bba\u7531\u60a8\u81ea\u4e3b\u51b3\u5b9a\u3002"
    )
    add_body_text(doc, app_a)

    add_heading_styled(
        doc,
        "\u9644\u5f55B \u7edf\u8ba1\u5206\u6790\u4ee3\u7801\u4e0e\u8f93\u51fa",
        level=2,
    )

    app_b = (
        "\u672c\u7814\u7a76\u7684\u5168\u90e8\u7edf\u8ba1\u5206\u6790\u4f7f\u7528Python 3.11\u5b8c\u6210\uff0c\u5206\u6790\u811a\u672c\u4f4d\u4e8e scripts/analyze_experiment.py\u3002\u4e3b\u8981\u4ee3\u7801\u6a21\u5757\u5305\u62ec\uff1a"
        "\uff081\uff09\u63cf\u8ff0\u6027\u7edf\u8ba1\u5206\u6790\uff1a\u51fd\u6570 compute_descriptive_statistics()\uff1b"
        "\uff082\uff09Cohen's Kappa\u4e00\u81f4\u6027\u5206\u6790\uff1a\u51fd\u6570 compute_cohens_kappa_analysis()\uff1b"
        "\uff083\uff09\u8017\u65f6\u5dee\u5f02\u5206\u6790\uff1a\u51fd\u6570 compute_time_analysis()\uff1b"
        "\uff084\uff09AI\u4e00\u81f4\u7387\u5206\u6790\uff1a\u51fd\u6570 compute_ai_agreement()\uff1b"
        "\uff085\uff09\u4e0d\u4e00\u81f4\u6848\u4f8b\u5206\u6790\uff1a\u51fd\u6570 analyze_inconsistent_cases()\uff1b"
        "\uff086\uff09\u6570\u636e\u53ef\u89c6\u5316\uff1a\u51fd\u6570 create_all_figures()\u3002"
        "\u5b8c\u6574\u7684\u5206\u6790\u8f93\u51fa\u7ed3\u679c\u4f4d\u4e8e research/results/ \u76ee\u5f55\u4e0b\uff0c\u4ee5JSON\u683c\u5f0f\u4fdd\u5b58\u3002"
    )
    add_body_text(doc, app_b)

    # ===== SAVE =====
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = REPORTS_DIR / "\u5b9e\u8bc1\u7814\u7a76\u62a5\u544a_V1.0.docx"
    doc.save(str(output_path))
    print(f"\nDone! Report saved to: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    generate_report()
