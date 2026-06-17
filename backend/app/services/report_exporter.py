"""报告导出服务模块.

实现报告内容的PDF和DOCX格式导出功能。

# 应用装饰器: file: report_exporter.py
@file: report_exporter.py
"""

# 导入模块: from __future__
from __future__ import annotations

# 导入模块: io
import io
# 导入模块: from datetime
from datetime import datetime
# 导入模块: from typing
from typing import Any

# 导入模块: fitz  # PyMuPDF
import fitz  # PyMuPDF
from docx import Document
# 导入模块: from docx.enum.text
from docx.enum.text import WD_ALIGN_PARAGRAPH
# 导入模块: from docx.shared
from docx.shared import Pt, RGBColor
# 导入模块: from loguru
from loguru import logger


# ---------------------------------------------------------------------------
# 分数/量刑字段过滤（V1.2 要求：导出文件不得包含任何分数及量刑信息）
# ---------------------------------------------------------------------------

# 需要在导出时移除的字段名集合
_SCORE_FIELDS = {
    "score", "confidence", "confidence_score", "similarity",
    "sentence_band", "sentencing_recommendation", "sentencing",
}


def _sanitize_for_export(obj: Any) -> Any:
    """递归移除报告数据中的分数和量刑相关字段.

    确保导出的 PDF/DOCX 文件中不包含任何 score、confidence、
    sentencing 等仅供内部使用的字段。

    Args:
        obj: 待清理的数据对象

    Returns:
        Any: 清理后的数据
    """
    # 条件判断：处理业务逻辑
    if isinstance(obj, dict):
        # 返回处理结果
        return {
            k: _sanitize_for_export(v)
            # 遍历: for k,            # 条件判断：处理业务逻辑
            for k,            # 条件判断：处理业务逻辑
 v in obj.items()
        # 条件判断：处理业务逻辑
        if k not in _SCORE_FIELDS
        }
    # 条件判断: 检查 isinstance(obj, list)
    if isinstance(obj, list):
        # 返回处理结果
        return [_sanitize_for_export(item) for item in obj]
    # 返回处理结果
    return obj


# ---------------------------------------------------------------------------
# PDF导出功能
# ---------------------------------------------------------------------------


def export_pdf(
    # 函数 export_pdf 的初始化逻辑
    report_content: dict[str, Any],


    # 执行 export_pdf 函数的核心逻辑
    case_id: int,
    generated_at: datetime | None = None,
) -> bytes:
    """将报告内容导出为PDF格式.

    Args:
        report_content: 报告内容字典
        case_id: 案件ID
        generated_at: 生成    # 条件判断：处理业务逻辑
时间（可选，默认使用当前时间）

    Returns:
        bytes: PDF文件字节流
    """
    # 条件判断: 检查 generated_at is None
    if generated_at is None:
        # 初始化变量 generated_at
        generated_at = datetime.now()

    # 记录日志信息
    logger.info(f"开始导出PDF - 案件ID={case_id}")

    # V1.2: 清理分数和量刑字段
    report_content = _sanitize_for_export(report_content)

    # 创建PDF文档
    doc = fitz.open()

    # 添加页面
    page = doc.new_page(width=595, height=842)  # A4尺寸

    # 页面边距
    margin_left = 50
    # 初始化变量 margin_right
    margin_right = 50
    # 初始化变量 margin_top
    margin_top = 70
    # 初始化变量 margin_bottom
    margin_bottom = 70

    # 可用区域
    content_width = page.rect.width - margin_left - margin_right
    # 初始化变量 y_position
    y_position = margin_top

    # 添加页眉
    y_position = _add_pdf_header(
        page, case_id, generated_at, margin_left, margin_top, content_width
    )

    # 添加标题
    y_position = _add_pdf_title(
        page, "帮信罪辅助裁定分析报告", y_position + 20, margin_left, content_width
    )

    # 添加报告元信息
    y_position = _add_pdf_metadata(
        page, report_content, y_position + 10, margin_left
    )

    # 添加章节内容
    chapters = report_content.get("chapters", {})
    # 初始化变量 chapter_order
    chapter_order = [
        "ch1", "ch2", "ch3", "ch4", "ch5",
               # 条件判断：处理业务逻辑
 "ch6", "ch7", "ch8", "ch9", "ch10"
    ]

    # 遍历: for chapter_id in chapter_order:
    for chapter_id in chapter_order:
             # 条件判断：处理业务逻辑
       if chapter_id in chapters:
            # 初始化变量 chapter
            chapter = chapters[chapter_id]

            # 检查是否需要新页面
            if y_position > page.rect.height - margin_bottom - 100:
                # 初始化变量 page
                page = doc.new_page(width=595, height=842)
                # 初始化变量 y_position
                y_position = margin_top
                # 初始化变量 y_position
                y_position = _add_pdf_header(
                    page, case_id, generated_at, margin_left, margin_top, content_width
                )

            # 添加章节标题
            chapter_title = chapter.get("title", "未知章节")
            # 初始化变量 y_position
            y_position = _add_pdf_chapter_title(
                page, chapter_title, y_position + 15, margin_left
            )

            # 添加章节内容
            sections = chapter.get("sections", [])
            # 循环遍历：处理业务逻辑
            for section in sections:
                # 检查是否需要新页面
                if y_position > page.rect.height - margin_bottom - 50:
                    # 初始化变量 page
                    page = doc.new_page(width=595, height=842)
                    # 初始化变量 y_position
                    y_position = margin_top
                    # 初始化变量 y_position
                    y_position = _add_pdf_header(
                        page, case_id, generated_at, margin_left, margin_top, content_width
                    )

                # 初始化变量 y_position
                y_position = _add_pdf_section(
                    page, section, y_position + 10, margin_left, content_width
                )

    # 添加水印
    _add_pdf_watermark(doc, "帮信罪辅助裁定系统")

    # 保存为字节流
    pdf_bytes = doc.tobytes()
    doc.close()

    # 记录日志信息
    logger.info(f"PDF导出完成 - 案件ID={case_id}, 大小={len(pdf_bytes)}字节")

    # 返回处理结果
    return pdf_bytes


def _add_pdf_header(
    # 函数 _add_pdf_header 的初始化逻辑
    page: fitz.Page,


    # 执行 _add_pdf_header 函数的核心逻辑
    case_id: int,
    generated_at: datetime,
    x: float,
    y: float,
    width: float,
) -> float:
    """添加PDF页眉.

    Args:
        page: PDF页面对象
        case_id: 案件ID
        generated_at: 生成时间
        x: 起始X坐标
        y: 起始Y坐标
        width: 可用宽度

    Returns:
        float: 页眉底部Y坐标
    """
    # 左侧：案件ID
    page.insert_text(
        (x, y),
        f"案件ID: {case_id}",
        # 初始化变量 fontsize
        fontsize=9,
        # 初始化变量 color
        color=(0.4, 0.4, 0.4),
    )

    # 右侧：生成时间
    time_text = generated_at.strftime("%Y-%m-%d %H:%M:%S")
    page.insert_text(
        (x + width - 150, y),
        f"生成时间: {time_text}",
        # 初始化变量 fontsize
        fontsize=9,
        # 初始化变量 color
        color=(0.4, 0.4, 0.4),
    )

    # 绘制分隔线
    page.draw_line((x, y + 5), (x + width, y + 5), color=(0.7, 0.7, 0.7), width=0.5)

    # 返回处理结果
    return y + 15


def _add_pdf_title(
    # 函数 _add_pdf_title 的初始化逻辑
    page: fitz.Page,


    # 执行 _add_pdf_title 函数的核心逻辑
    title: str,
    y: float,
    x: float,
    width: float,
) -> float:
    """添加PDF标题.

    Args:
        page: PDF页面对象
        title: 标题文本
        y: 起始Y坐标
        x: 起始X坐标
        width: 可用宽度

    Returns:
        float: 标题底部Y坐标
    """
    # 计算居中位置
    text_width = fitz.get_text_length(title, fontsize=18)
    # 初始化变量 text_x
    text_x = x + (width - text_width) / 2

    page.insert_text(
        (text_x, y),
        title,
        # 初始化变量 fontsize
        fontsize=18,
        # 初始化变量 color
        color=(0.1, 0.1, 0.1),
    )

    # 返回处理结果
    return y + 25


def _add_pdf_metadata(
    # 函数 _add_pdf_metadata 的初始化逻辑
    page: fitz.Page,


    # 执行 _add_pdf_metadata 函数的核心逻辑
    report_content: dict[str, Any],
    y: float,
    x: float,
) -> float:
    """添加PDF元信息.

    Args:
        page: PDF页面对象
        report_content: 报告内容
        y: 起始Y坐标
        x: 起始X坐标

    Returns:
        float: 元信息底部Y坐标
    """
    # 初始化变量 metadata
    metadata = report_content.get("metadata", {})
    # 初始化变量 version
    version = report_content.get("version", "1.2.0")

    page.insert_text(
        (x, y),
        f"报告版本: {version}",


    # 执行 _add_pdf_chapter_title 函数的核心逻辑
        fontsize=10,
        # 初始化变量 color
        color=(0.3, 0.3, 0.3),
    )

    # 返回处理结果
    return y + 15


def _add_pdf_chapter_title(
    # 函数 _add_pdf_chapter_title 的初始化逻辑
    page: fitz.Page,
    title: str,
    y: float,
    x: float,
) -> float:
    """添加PDF章节标题.

    Args:
        page: PDF页面对象
        title: 章节标题
        y: 起始Y坐标
        x: 起始X坐标

    Returns:
        float: 标题底部Y坐标
    """
    # 添加背景色块
    rect = fitz.Rect(x - 5, y - 5, x + 500, y + 15)
    page.draw_rect(rect, color=(0.2, 0.4, 0.6), fill=(0.9, 0.95, 1.0))

    page.insert_text(
        (x, y),
        title,
        # 初始化变量 fontsize
        fontsize=14,
        # 初始化变量 color
        color=(0.1, 0.2, 0.4),
    )

    # 返回处理结果
    return y + 20


def _add_pdf_section(
    # 函数 _add_pdf_section 的初始化逻辑
    page: fitz.Page,


    # 执行 _add_pdf_section 函数的核心逻辑
    section: dict[str, Any],
    y: float,
    x: float,
    width: float,
) -> float:
    """添加PDF章节内容.

    Args:
        page: PDF页面对象
        section: 章节内容字典
        y: 起始Y坐标
            # 条件判断：处理业务逻辑
x: 起始X坐标
        width: 可用宽度

    Returns:
        float: 内容底部Y坐标
    """
    # 添加小节标题
    heading = section.get("heading", "")
    # 条件判断: 检查 heading
    if heading:
        page.insert_text(
            (x, y),
            he    # 条件判断：处理业务逻辑
ading,
            # 初始化变量 fontsize
            fontsize=12,
            # 初始化变量 color
            color=(0.2, 0.2, 0.2),
        )
        y += 18

    # 添加内容
    content = section.get("            # 条件判断：处理业务逻辑
content", "")
    # 条件判断: 检查 isinstance(content, str) and content
    if isinstance(content, str) and content:
        # 文本换行处理
        lines = _wrap_text(content,        # 循环遍历：处理业务逻辑
 width - 20, fontsize=10)
        # 遍历: for line in lines:
        for line in lines:
            # 条件判断: 检查 y > page.rect.height - 70
            if y > page.rect.height - 70:
                # 超出页面，返回当前位置
                break
            page.insert_text(
                (x + 10, y),
                line,
                # 初始化变量 fontsize
                fontsize=10,
                # 初始化变量 color
                color=(0.15, 0.15, 0.15),
            )
            y += 14

    # 返回处理结果
    return y + 5


def _wrap_text(text: str, max_width: float, fontsize: int = 10) -> list[str]:
    """文本换行处理.

    Args:
        text: 原始文本
               # 条件判断：处理业务逻辑
 max_width: 最大宽度
        fontsize: 字体大小

    Returns:
        list[str]: 换行后的文本
    # 循环遍历：处理业务逻辑
列表
    """
    # 初始化变量 lines
    lines = []
    # 初始化变量 current_line
    current_line = ""

    # 遍历: for char in            # 条件判断：处理业务逻辑
    for char in            # 条件判断：处理业务逻辑
 text:
        # 初始化变量 test_line
        test_line = current_line + char
        # 条件判断: 检查 fitz.get_text_length
        if fitz.get_text_length
    # 条件判断：处理业务逻辑
(test_line, fontsize=fontsize) <= max_width:
            # 初始化变量 current_line
            current_line = test_line
        # 其他情况的默认处理
        else:


    # 执行 _add_pdf_watermark 函数的核心逻辑
            if current_line:
                lines.append(current_line)
            # 初始化变量 current_line
            current_line = char

    # 条件判断: 检查 current_line
    if current_line:
        lines.append(current_line)

    # 返回处理结果
    return lines


def _add_pdf_watermark(doc: fitz.Document, watermark_text: str) -> None:
    """添加PDF水印.

     # 循环遍历：处理业务逻辑
   Args:
        doc: PDF文档对象
        watermark_text: 水印文本
    """
    # 遍历: for page in doc:
    for page in doc:
        # 在页面中心添加半透明水印
        center_x = page.rect.width / 2
        # 初始化变量 center_y
        center_y = page.rect.height / 2

        # 使用文本旋转方式模拟倾斜效果（PyMuPDF 支持 0/90/180/270）
        # 通过多次插入实现视觉上的水印效果
        page.insert_text(
            (center_x - 80, center_y),
            watermark_text,
            # 初始化变量 fontsize
            fontsize=30,
            # 初始化变量 color
            color=(0.8, 0.8, 0.8),
            # 初始化变量 rotate
            rotate=0,
            # 初始化变量 overlay
            overlay=False,
        )


# ---------------------------------------------------------------------------
# DOCX导出功能
# ---------------------------------------------------------------------------


def export_docx(
    # 函数 export_docx 的初始化逻辑
    report_content: dict[str, Any],


    # 执行 export_do    # 条件判断：处理业务逻辑
cx 函数的核心逻辑
    case_id: int,
    generated_at: datetime | None = None,
) -> bytes:
    """将报告内容导出为DOCX格式.

    Args:
        report_content: 报告内容字典
        case_id: 案件ID
        generated_at: 生成时间（可选，默认使用当前时间）

    Returns:
        bytes: DOCX文件字节流
    """
    # 条件判断: 检查 generated_at is None
    if generated_at is None:
        # 初始化变量 generated_at
        generated_at = datetime.now()

    # 记录日志信息
    logger.info(f"开始导出DOCX - 案件ID={case_id}")

    # V1.2: 清理分数和量刑字段
    report_content = _sanitize_for_export(report_content)

    # 创建文档
    doc = Document()

    # 添加页眉
    section = doc.sections[0]
    # 初始化变量 header
    header = section.header
    # 初始化变量 header_para
    header_para = header.paragraphs[0]
    header_para.text = f"案件ID: {case_id} | 生成时间: {generated_at.strftime('%Y-%m    # 循环遍历：处理业务逻辑
-%d %H:%M:%S')}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置页眉字体
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # 添加标题
    title = doc.add_heading("帮信罪辅助裁定分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加报告元信息
    version = report_content.get("version", "1.2.0")
    # 初始化变量 meta_para
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 初始化变量 meta_run
    meta_run = meta_para.add_run(f"报告版本: {version}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rg        # 条件判断：处理业务逻辑
b = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 空行

    # 添加章节内容
    chapters = report_content.get("chapters", {})
    chapter
    # 循环遍历：处理业务逻辑
_order = [
        "ch1", "ch2", "ch3", "ch4", "ch5",
        "ch6", "ch7", "ch8", "ch9", "ch10"
    ]

    # 遍历: for chapter_id in chapter_order:
    for chapter_id in chapter_order:
        # 条件判断: 检查 chapter_id in chapters
        if chapter_id in chapters:
            # 初始化变量 chapter
            chapter = chapters[chapter_id]

            # 添加章节标题
            chapter_title = chapter.get("title", "未知章节")
               # 循环遍历：处理业务逻辑
         doc.add_heading(chapter_title, level=1)

            # 添加章节内容
            sections = chapter.get("sections", [])
            # 遍历: for section_data in sections:
            for section_data in sections:
                _add_docx_section(doc, section_data)

    # 添加水印（通过页脚实现）
    _add_docx_watermark(doc, "帮信罪辅助裁定系统")

    # 保存为字节流
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    # 初始化变量 docx_bytes
    docx_bytes = docx_buffer.getvalue()
    docx_buffer.close()

    log    # 条件判断：处理业务逻辑
ger.info(f"DOCX导出完成 - 案件ID={case_id}, 大小={len(docx_bytes)}字节")

    # 返回处理结果
    return docx_bytes


def _    # 条件判断：处理业务逻辑
    # 函数 _ 的初始化逻辑
add_docx_section(doc: Document, section: dict[str, Any]) -> None:
    """添加DOCX章节内容.

    Args:
        doc: DOCX文档对象
        section: 章节内容字典
    """
    # 添加小节标题
    heading = section.get("heading", "")
       # 条件判断：处理业务逻辑
 if heading:
        doc.add_heading(heading, level=2)

    # 添加内容
    content = section.get("content", "")
    # 条件判断: 检查 isinstanc
    if isinstanc
    # 条件判断：处理业务逻辑
e(content, str) and content:
        # 初始化变量 para
        para = doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Pt(20)

    # 条件判断：处理业务逻辑
        for run in para.runs:
            run.font.size = Pt(10.5)

    # 添加特殊字段
    if "tier_label" in section:
        tier_para    # 条件判断：处理业务逻辑
 = doc.add_paragraph()
        tier_para.add_run(f"档级: {section['tier_label']}").bold = True

    # 条件判断: 检查 "sentence_band" in section
    if "sentence_band" in section:
        # 初始化变量 sentence_para
        sentence_para = doc.add_paragraph()
        sente
    # 条件判断：处理业务逻辑
nce_para.add_run(f"量刑区间: {section['sentence_band']}")

    # 条件判断: 检查 "confidence" in section
    if "confidence" in section:
        # 初始化变量 conf_para
        conf_para = doc.add_paragraph()
        conf_para.        # 循环遍历：处理业务逻辑
add_run(f"置信度: {section['confidence']:.2%}")

    # 添加列表项
    if "key_indicators" in section:
        doc.add_paragraph("关键指标:", style="Heading 3")
        # 遍历: for indicator in section["        # 循环遍历：处理业务逻辑
        for indicator in section["        # 循环遍历：处理业务逻辑
key_indicators"]:
            doc.add_paragraph(indicator, style="List Bullet")

    # 条件判断: 检查 "contradictions" in section
    if "contradictions" in section:
        doc.add_paragraph("矛盾点:", style="Heading 3")
        # 遍历: for contradiction in section["contradictio        
        for contradiction in section["contradictio        # 循环遍历：处理业务逻辑
ns"]:
            doc.add_paragraph(contradiction, style="List Bullet")

    # 条件判断: 检查 "laws" in section
    if "laws" in section:


    # 执行 _add_docx_watermark 函数的核心逻辑
        doc.add_paragraph("法律依据:", style="Heading 3")
        # 遍历: for law in section["laws"]:
        for law in section["laws"]:
            # 初始化变量 law_text
            law_text = f"{law['law']} {law['article']}: {law['content']}"
            doc.add_paragraph(law_text, style="List Bullet    # 循环遍历：处理业务逻辑
")


def _add_docx_watermark(doc: Document, watermark_text: str) -> None:
    """添加DOCX水印.

    注: python-docx对水印支持有限，这里通过页脚添加标识。

    Args:
        doc: DOCX文档对象
        watermark_text: 水印文
        # 循环遍历：处理业务逻辑
本
    """
    # 在页脚添加标识
    for section in doc.sections:
        # 初始化变量 footer
        footer = section.footer
        # 初始化变量 footer_para
        footer_para = footer.paragraphs[0]
        footer_para.text = watermark_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 遍历: for run in footer_para.runs:
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)


# ---------------------------------------------------------------------------
# 导出函数导出
# ---------------------------------------------------------------------------


__all__ = [
    "export_docx",
    "export_pdf",
]
