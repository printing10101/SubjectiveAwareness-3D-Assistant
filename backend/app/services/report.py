"""报告服务统一模块.

整合报告生成、查询和导出功能，提供完整的报告管理能力。
包含：
- 报告内容生成（10章结构化报告）
- 报告查询和分页
- PDF/DOCX 格式导出

@file: report.py
"""

from __future__ import annotations

import io
from datetime import datetime
from math import ceil
from typing import Any

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from loguru import logger
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.analysis import Analysis
from app.models.case import Case
from app.models.report import Report, ReportReview
from app.types.analysis import AnalysisReport
from app.types.analysis_v2 import AnalysisResultV2, TierEnum


# ---------------------------------------------------------------------------
# 引用段落数据结构
# ---------------------------------------------------------------------------


class Citation:
    """引用段落数据结构.

    指向案件原文的字符偏移量。

    Attributes:
        start: 起始字符位置
        end: 结束字符位置
        text: 引用文本内容
    """

    def __init__(self, start: int, end: int, text: str) -> None:
        """初始化引用段落.

        Args:
            start: 起始字符位置
            end: 结束字符位置
            text: 引用文本内容
        """
        self.start = start
        self.end = end
        self.text = text

    def to_dict(self) -> dict[str, Any]:
        """转换为字典格式."""
        return {
            "start": self.start,
            "end": self.end,
            "text": self.text,
        }


# ---------------------------------------------------------------------------
# 报告查询服务
# ---------------------------------------------------------------------------


async def list_reports(
    db: AsyncSession,
    page: int = 1,
    page_size: int = 20,
) -> dict:
    """分页查询分析报告列表.

    Args:
        db: 异步数据库会话
        page: 页码（从 1 开始）
        page_size: 每页条数

    Returns:
        dict: 包含 total、page、page_size、total_pages、reports 的字典
    """
    page = max(1, page)
    page_size = min(max(1, page_size), 100)

    count_result = await db.execute(select(func.count(Analysis.id)))
    total: int = count_result.scalar() or 0

    offset = (page - 1) * page_size
    result = await db.execute(
        select(Analysis)
        .order_by(Analysis.created_at.desc())
        .offset(offset)
        .limit(page_size)
    )
    analyses = list(result.scalars().all())

    reports = [_format_analysis(a) for a in analyses]

    return {
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": max(1, ceil(total / page_size)) if total > 0 else 0,
        "reports": reports,
    }


def _format_analysis(a: Analysis) -> AnalysisReport:
    """格式化单条分析记录.

    Args:
        a: 分析记录实例

    Returns:
        AnalysisReport: 格式化后的字典
    """
    return {
        "id": a.id if a.id is not None else 0,
        "case_id": int(a.case_id) if a.case_id is not None else None,
        "knowledge_score": (
            float(a.knowledge_score) if a.knowledge_score is not None else None
        ),
        "mode": str(a.mode),
        "result": str(a.result_json),
        "created_at": a.created_at.isoformat() if a.created_at else None,
    }


# ---------------------------------------------------------------------------
# 报告内容生成
# ---------------------------------------------------------------------------


def ch1_basic_info(
    case: Case,
    analysis_result: AnalysisResultV2,
    generated_at: datetime,
) -> dict[str, Any]:
    """生成基本信息章节.

    包含案件编号、案件名称、分析日期等元数据。

    Args:
        case: 案件对象
        analysis_result: 分析结果
        generated_at: 报告生成时间

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第1章：基本信息 - 案件ID={case.id}")

    return {
        "chapter_id": "ch1",
        "title": "基本信息",
        "sections": [
            {
                "heading": "案件基本信息",
                "content": {
                    "案件编号": str(case.id),
                    "案件名称": case.title,
                    "案件状态": case.status.value if case.status else "未知",
                    "分析日期": generated_at.strftime("%Y年%m月%d日 %H:%M:%S"),
                    "报告版本": "1.0.0",
                },
                "citations": [],
            }
        ],
    }


def ch2_fact_summary(
    case: Case,
    analysis_result: AnalysisResultV2,
) -> dict[str, Any]:
    """生成事实摘要章节.

    提炼案件核心事实要素。

    Args:
        case: 案件对象
        analysis_result: 分析结果

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第2章：事实摘要 - 案件ID={case.id}")

    fact_text = case.description or "无案件描述"
    citations = []
    if case.description:
        citations.append(
            Citation(0, min(len(case.description), 200), case.description[:200]).to_dict()
        )

    return {
        "chapter_id": "ch2",
        "title": "事实摘要",
        "sections": [
            {
                "heading": "案件事实概述",
                "content": fact_text[:500] if len(fact_text) > 500 else fact_text,
                "citations": citations,
            }
        ],
    }


def ch3_dimensional_analysis(
    case: Case,
    analysis_result: AnalysisResultV2,
) -> dict[str, Any]:
    """生成维度分析章节.

    从多个维度对案件进行剖析。

    Args:
        case: 案件对象
        analysis_result: 分析结果

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第3章：维度分析 - 案件ID={case.id}")

    sections = []

    if "dimension1" in analysis_result:
        dim1 = analysis_result["dimension1"]
        sections.append({
            "heading": "维度1：构成要件分析",
            "content": dim1.get("reasoning", "无分析内容"),
            "tier": dim1.get("tier", "T2"),
            "confidence": dim1.get("confidence", 0.0),
            "key_indicators": dim1.get("key_indicators", []),
            "citations": [],
        })

    if "dimension2" in analysis_result:
        dim2 = analysis_result["dimension2"]
        sections.append({
            "heading": "维度2：情节模式分析",
            "content": dim2.get("reasoning", "无分析内容"),
            "tier": dim2.get("tier", "T2"),
            "confidence": dim2.get("confidence", 0.0),
            "pattern_match": dim2.get("pattern_match", ""),
            "citations": [],
        })

    if "dimension3" in analysis_result:
        dim3 = analysis_result["dimension3"]
        sections.append({
            "heading": "维度3：矛盾分析",
            "content": dim3.get("reasoning", "无分析内容"),
            "tier": dim3.get("tier", "T2"),
            "confidence": dim3.get("confidence", 0.0),
            "contradictions": dim3.get("contradictions", []),
            "citations": [],
        })

    return {
        "chapter_id": "ch3",
        "title": "维度分析",
        "sections": sections,
    }


def ch4_triggered_rules(
    case: Case,
    analysis_result: AnalysisResultV2,
    rule_hits: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """生成触发规则章节.

    列出所有命中的规则及其详情。

    Args:
        case: 案件对象
        analysis_result: 分析结果
        rule_hits: 规则命中列表

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第4章：触发规则 - 案件ID={case.id}")

    triggered_rule_ids = analysis_result.get("triggered_rule_ids", [])

    sections = []
    if rule_hits:
        for rule in rule_hits:
            sections.append({
                "heading": f"规则 {rule.get('rule_id', '未知')}",
                "content": rule.get("description", "无描述"),
                "rule_id": rule.get("rule_id"),
                "severity": rule.get("severity"),
                "citations": [],
            })
    elif triggered_rule_ids:
        for rule_id in triggered_rule_ids:
            sections.append({
                "heading": f"规则 {rule_id}",
                "content": f"命中规则：{rule_id}",
                "rule_id": rule_id,
                "citations": [],
            })

    return {
        "chapter_id": "ch4",
        "title": "触发规则",
        "sections": sections or [{"heading": "无触发规则", "content": "本次分析未触发任何规则", "citations": []}],
    }


def ch5_fact_tags(
    case: Case,
    analysis_result: AnalysisResultV2,
    tags: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """生成事实标签章节.

    展示案件相关标签及分类。

    Args:
        case: 案件对象
        analysis_result: 分析结果
        tags: 标签列表

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第5章：事实标签 - 案件ID={case.id}")

    matched_tag_ids = analysis_result.get("matched_tag_ids", [])

    sections = []
    if tags:
        for tag in tags:
            sections.append({
                "heading": tag.get("name", "未知标签"),
                "content": tag.get("description", "无描述"),
                "tag_id": tag.get("tag_id"),
                "category": tag.get("category"),
                "citations": [],
            })
    elif matched_tag_ids:
        for tag_id in matched_tag_ids:
            sections.append({
                "heading": f"标签 {tag_id}",
                "content": f"命中标签：{tag_id}",
                "tag_id": tag_id,
                "citations": [],
            })

    return {
        "chapter_id": "ch5",
        "title": "事实标签",
        "sections": sections or [{"heading": "无事实标签", "content": "本次分析未匹配任何标签", "citations": []}],
    }


def ch6_conflict_results(
    case: Case,
    analysis_result: AnalysisResultV2,
) -> dict[str, Any]:
    """生成冲突结果章节.

    分析案件中的矛盾点及处理结果。

    Args:
        case: 案件对象
        analysis_result: 分析结果

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第6章：冲突结果 - 案件ID={case.id}")

    conflicts = analysis_result.get("conflicts", [])

    sections = []
    if conflicts:
        for conflict in conflicts:
            sections.append({
                "heading": conflict.get("type", "冲突"),
                "content": conflict.get("description", "无描述"),
                "severity": conflict.get("severity"),
                "resolution": conflict.get("resolution", "待处理"),
                "citations": [],
            })

    return {
        "chapter_id": "ch6",
        "title": "冲突结果",
        "sections": sections or [{"heading": "无冲突", "content": "本次分析未发现冲突项", "citations": []}],
    }


def ch7_similar_cases(
    case: Case,
    analysis_result: AnalysisResultV2,
    similar_cases: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """生成相似案例章节.

    展示参考案例及其相似度分析。

    Args:
        case: 案件对象
        analysis_result: 分析结果
        similar_cases: 相似案例列表

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第7章：相似案例 - 案件ID={case.id}")

    sections = []
    if similar_cases:
        for similar in similar_cases:
            sections.append({
                "heading": similar.get("title", "未知案例"),
                "content": similar.get("summary", "无摘要"),
                "case_id": similar.get("case_id"),
                "similarity": similar.get("similarity", 0.0),
                "verdict": similar.get("verdict"),
                "citations": [],
            })

    return {
        "chapter_id": "ch7",
        "title": "相似案例",
        "sections": sections or [{"heading": "无相似案例", "content": "未找到相似案例", "citations": []}],
    }


def ch8_sentencing(
    case: Case,
    analysis_result: AnalysisResultV2,
) -> dict[str, Any]:
    """生成量刑建议章节.

    提供基于分析的量刑参考。

    Args:
        case: 案件对象
        analysis_result: 分析结果

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第8章：量刑建议 - 案件ID={case.id}")

    final_verdict = analysis_result.get("final_verdict", {})
    tier_str = final_verdict.get("final_tier", "T2")

    try:
        tier = TierEnum(tier_str)
        sentence_band = tier.sentence_band
        tier_label = tier.chinese_label
    except ValueError:
        tier = TierEnum.T2
        tier_str = tier.value
        sentence_band = tier.sentence_band
        tier_label = tier.chinese_label

    sections = [
        {
            "heading": "量刑建议",
            "content": f"根据综合分析，本案建议量刑区间为：{sentence_band}",
            "tier": tier_str,
            "tier_label": tier_label,
            "sentence_band": sentence_band,
            "confidence": final_verdict.get("confidence", 0.0),
            "citations": [],
        }
    ]

    if "subjective_knowledge" in analysis_result:
        sections.append({
            "heading": "主观明知程度",
            "content": analysis_result["subjective_knowledge"],
            "citations": [],
        })

    if "sentence" in analysis_result:
        sections.append({
            "heading": "量刑建议详情",
            "content": analysis_result["sentence"],
            "citations": [],
        })

    return {
        "chapter_id": "ch8",
        "title": "量刑建议",
        "sections": sections,
    }


def ch9_legal_basis(
    case: Case,
    analysis_result: AnalysisResultV2,
) -> dict[str, Any]:
    """生成法律依据章节.

    列出案件涉及的法律法规。

    Args:
        case: 案件对象
        analysis_result: 分析结果

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第9章：法律依据 - 案件ID={case.id}")

    legal_basis = [
        {
            "law": "《中华人民共和国刑法》",
            "article": "第二百八十七条之二",
            "content": "帮助信息网络犯罪活动罪",
        },
        {
            "law": "《中华人民共和国刑法》",
            "article": "第二十五条",
            "content": "共同犯罪",
        },
    ]

    sections = [
        {
            "heading": "主要法律依据",
            "content": "本案涉及的主要法律法规",
            "laws": legal_basis,
            "citations": [],
        }
    ]

    return {
        "chapter_id": "ch9",
        "title": "法律依据",
        "sections": sections,
    }


def ch10_review_conclusion(
    case: Case,
    analysis_result: AnalysisResultV2,
) -> dict[str, Any]:
    """生成审查结论章节.

    总结分析结果及建议。

    Args:
        case: 案件对象
        analysis_result: 分析结果

    Returns:
        dict: 章节内容字典
    """
    logger.info(f"生成第10章：审查结论 - 案件ID={case.id}")

    final_verdict = analysis_result.get("final_verdict", {})
    tier_str = final_verdict.get("final_tier", "T2")

    try:
        tier = TierEnum(tier_str)
        tier_label = tier.chinese_label
    except ValueError:
        tier = TierEnum.T2
        tier_str = tier.value
        tier_label = tier.chinese_label

    conclusion_text = (
        f"经过对案件事实、证据及法律适用的综合分析，"
        f"本案建议认定为{tier_label}。"
        f"综合置信度为{final_verdict.get('confidence', 0.0):.2%}。"
    )

    sections = [
        {
            "heading": "审查结论",
            "content": conclusion_text,
            "final_tier": tier_str,
            "final_label": tier_label,
            "confidence": final_verdict.get("confidence", 0.0),
            "citations": [],
        }
    ]

    if "disclaimer" in analysis_result:
        sections.append({
            "heading": "免责声明",
            "content": analysis_result["disclaimer"],
            "citations": [],
        })

    return {
        "chapter_id": "ch10",
        "title": "审查结论",
        "sections": sections,
    }


def generate_report(
    analysis_result: AnalysisResultV2,
    case: Case,
    rule_hits: list[dict[str, Any]] | None = None,
    tags: list[dict[str, Any]] | None = None,
    similar_cases: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """生成完整报告内容.

    接收分析结果、案件信息、规则命中情况、标签及相似案例数据，
    返回包含10个章节的结构化报告内容。

    Args:
        analysis_result: 分析结果（V2格式）
        case: 案件对象
        rule_hits: 规则命中列表
        tags: 标签列表
        similar_cases: 相似案例列表

    Returns:
        dict: 包含10个章节的完整报告内容
    """
    logger.info(f"开始生成报告 - 案件ID={case.id}")

    generated_at = datetime.now()

    chapters = [
        ch1_basic_info(case, analysis_result, generated_at),
        ch2_fact_summary(case, analysis_result),
        ch3_dimensional_analysis(case, analysis_result),
        ch4_triggered_rules(case, analysis_result, rule_hits),
        ch5_fact_tags(case, analysis_result, tags),
        ch6_conflict_results(case, analysis_result),
        ch7_similar_cases(case, analysis_result, similar_cases),
        ch8_sentencing(case, analysis_result),
        ch9_legal_basis(case, analysis_result),
        ch10_review_conclusion(case, analysis_result),
    ]

    report_content = {
        "report_id": None,
        "case_id": case.id,
        "generated_at": generated_at.isoformat(),
        "version": "1.0.0",
        "chapters": {ch["chapter_id"]: ch for ch in chapters},
        "metadata": {
            "total_chapters": len(chapters),
            "analysis_timestamp": analysis_result.get("timestamp"),
            "fallback": analysis_result.get("fallback", False),
        },
    }

    logger.info(f"报告生成完成 - 共{len(chapters)}章")

    return report_content


# ---------------------------------------------------------------------------
# PDF 导出功能
# ---------------------------------------------------------------------------


def export_pdf(
    report_content: dict[str, Any],
    case_id: int,
    generated_at: datetime | None = None,
) -> bytes:
    """将报告内容导出为PDF格式.

    Args:
        report_content: 报告内容字典
        case_id: 案件ID
        generated_at: 生成时间（可选，默认使用当前时间）

    Returns:
        bytes: PDF文件字节流
    """
    if generated_at is None:
        generated_at = datetime.now()

    logger.info(f"开始导出PDF - 案件ID={case_id}")

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)

    margin_left = 50
    margin_right = 50
    margin_top = 70
    margin_bottom = 70

    content_width = page.rect.width - margin_left - margin_right
    y_position = margin_top

    y_position = _add_pdf_header(
        page, case_id, generated_at, margin_left, margin_top, content_width
    )

    y_position = _add_pdf_title(
        page, "帮信罪辅助裁定分析报告", y_position + 20, margin_left, content_width
    )

    y_position = _add_pdf_metadata(
        page, report_content, y_position + 10, margin_left
    )

    chapters = report_content.get("chapters", {})
    chapter_order = [
        "ch1", "ch2", "ch3", "ch4", "ch5",
        "ch6", "ch7", "ch8", "ch9", "ch10"
    ]

    for chapter_id in chapter_order:
        if chapter_id in chapters:
            chapter = chapters[chapter_id]

            if y_position > page.rect.height - margin_bottom - 100:
                page = doc.new_page(width=595, height=842)
                y_position = margin_top
                y_position = _add_pdf_header(
                    page, case_id, generated_at, margin_left, margin_top, content_width
                )

            chapter_title = chapter.get("title", "未知章节")
            y_position = _add_pdf_chapter_title(
                page, chapter_title, y_position + 15, margin_left
            )

            sections = chapter.get("sections", [])
            for section in sections:
                if y_position > page.rect.height - margin_bottom - 50:
                    page = doc.new_page(width=595, height=842)
                    y_position = margin_top
                    y_position = _add_pdf_header(
                        page, case_id, generated_at, margin_left, margin_top, content_width
                    )

                y_position = _add_pdf_section(
                    page, section, y_position + 10, margin_left, content_width
                )

    _add_pdf_watermark(doc, "帮信罪辅助裁定系统")

    pdf_bytes = doc.tobytes()
    doc.close()

    logger.info(f"PDF导出完成 - 案件ID={case_id}, 大小={len(pdf_bytes)}字节")

    return pdf_bytes


def _add_pdf_header(
    page: fitz.Page,
    case_id: int,
    generated_at: datetime,
    x: float,
    y: float,
    width: float,
) -> float:
    """添加PDF页眉."""
    page.insert_text(
        (x, y),
        f"案件ID: {case_id}",
        fontsize=9,
        color=(0.4, 0.4, 0.4),
    )

    time_text = generated_at.strftime("%Y-%m-%d %H:%M:%S")
    page.insert_text(
        (x + width - 150, y),
        f"生成时间: {time_text}",
        fontsize=9,
        color=(0.4, 0.4, 0.4),
    )

    page.draw_line((x, y + 5), (x + width, y + 5), color=(0.7, 0.7, 0.7), width=0.5)

    return y + 15


def _add_pdf_title(
    page: fitz.Page,
    title: str,
    y: float,
    x: float,
    width: float,
) -> float:
    """添加PDF标题."""
    text_width = fitz.get_text_length(title, fontsize=18)
    text_x = x + (width - text_width) / 2

    page.insert_text(
        (text_x, y),
        title,
        fontsize=18,
        color=(0.1, 0.1, 0.1),
    )

    return y + 25


def _add_pdf_metadata(
    page: fitz.Page,
    report_content: dict[str, Any],
    y: float,
    x: float,
) -> float:
    """添加PDF元信息."""
    version = report_content.get("version", "1.0.0")

    page.insert_text(
        (x, y),
        f"报告版本: {version}",
        fontsize=10,
        color=(0.3, 0.3, 0.3),
    )

    return y + 15


def _add_pdf_chapter_title(
    page: fitz.Page,
    title: str,
    y: float,
    x: float,
) -> float:
    """添加PDF章节标题."""
    rect = fitz.Rect(x - 5, y - 5, x + 500, y + 15)
    page.draw_rect(rect, color=(0.2, 0.4, 0.6), fill=(0.9, 0.95, 1.0))

    page.insert_text(
        (x, y),
        title,
        fontsize=14,
        color=(0.1, 0.2, 0.4),
    )

    return y + 20


def _add_pdf_section(
    page: fitz.Page,
    section: dict[str, Any],
    y: float,
    x: float,
    width: float,
) -> float:
    """添加PDF章节内容."""
    heading = section.get("heading", "")
    if heading:
        page.insert_text(
            (x, y),
            heading,
            fontsize=12,
            color=(0.2, 0.2, 0.2),
        )
        y += 18

    content = section.get("content", "")
    if isinstance(content, str) and content:
        lines = _wrap_text(content, width - 20, fontsize=10)
        for line in lines:
            if y > page.rect.height - 70:
                break
            page.insert_text(
                (x + 10, y),
                line,
                fontsize=10,
                color=(0.15, 0.15, 0.15),
            )
            y += 14

    return y + 5


def _wrap_text(text: str, max_width: float, fontsize: int = 10) -> list[str]:
    """文本换行处理."""
    lines = []
    current_line = ""

    for char in text:
        test_line = current_line + char
        if fitz.get_text_length(test_line, fontsize=fontsize) <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = char

    if current_line:
        lines.append(current_line)

    return lines


def _add_pdf_watermark(doc: fitz.Document, watermark_text: str) -> None:
    """添加PDF水印."""
    for page in doc:
        center_x = page.rect.width / 2
        center_y = page.rect.height / 2

        page.insert_text(
            (center_x - 100, center_y),
            watermark_text,
            fontsize=30,
            color=(0.8, 0.8, 0.8),
            rotate=45,
            overlay=False,
        )


# ---------------------------------------------------------------------------
# DOCX 导出功能
# ---------------------------------------------------------------------------


def export_docx(
    report_content: dict[str, Any],
    case_id: int,
    generated_at: datetime | None = None,
) -> bytes:
    """将报告内容导出为DOCX格式.

    Args:
        report_content: 报告内容字典
        case_id: 案件ID
        generated_at: 生成时间（可选，默认使用当前时间）

    Returns:
        bytes: DOCX文件字节流
    """
    if generated_at is None:
        generated_at = datetime.now()

    logger.info(f"开始导出DOCX - 案件ID={case_id}")

    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"案件ID: {case_id} | 生成时间: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    title = doc.add_heading("帮信罪辅助裁定分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    version = report_content.get("version", "1.0.0")
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"报告版本: {version}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    chapters = report_content.get("chapters", {})
    chapter_order = [
        "ch1", "ch2", "ch3", "ch4", "ch5",
        "ch6", "ch7", "ch8", "ch9", "ch10"
    ]

    for chapter_id in chapter_order:
        if chapter_id in chapters:
            chapter = chapters[chapter_id]

            chapter_title = chapter.get("title", "未知章节")
            doc.add_heading(chapter_title, level=1)

            sections = chapter.get("sections", [])
            for section_data in sections:
                _add_docx_section(doc, section_data)

    _add_docx_watermark(doc, "帮信罪辅助裁定系统")

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()
    docx_buffer.close()

    logger.info(f"DOCX导出完成 - 案件ID={case_id}, 大小={len(docx_bytes)}字节")

    return docx_bytes


def _add_docx_section(doc: Document, section: dict[str, Any]) -> None:
    """添加DOCX章节内容."""
    heading = section.get("heading", "")
    if heading:
        doc.add_heading(heading, level=2)

    content = section.get("content", "")
    if isinstance(content, str) and content:
        para = doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Pt(20)
        for run in para.runs:
            run.font.size = Pt(10.5)

    if "tier_label" in section:
        tier_para = doc.add_paragraph()
        tier_para.add_run(f"档级: {section['tier_label']}").bold = True

    if "sentence_band" in section:
        sentence_para = doc.add_paragraph()
        sentence_para.add_run(f"量刑区间: {section['sentence_band']}")

    if "confidence" in section:
        conf_para = doc.add_paragraph()
        conf_para.add_run(f"置信度: {section['confidence']:.2%}")

    if "key_indicators" in section:
        doc.add_paragraph("关键指标:", style="Heading 3")
        for indicator in section["key_indicators"]:
            doc.add_paragraph(indicator, style="List Bullet")

    if "contradictions" in section:
        doc.add_paragraph("矛盾点:", style="Heading 3")
        for contradiction in section["contradictions"]:
            doc.add_paragraph(contradiction, style="List Bullet")

    if "laws" in section:
        doc.add_paragraph("法律依据:", style="Heading 3")
        for law in section["laws"]:
            law_text = f"{law['law']} {law['article']}: {law['content']}"
            doc.add_paragraph(law_text, style="List Bullet")


def _add_docx_watermark(doc: Document, watermark_text: str) -> None:
    """添加DOCX水印."""
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = watermark_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)


# ---------------------------------------------------------------------------
# 模块导出
# ---------------------------------------------------------------------------


__all__ = [
    "Citation",
    "list_reports",
    "generate_report",
    "export_pdf",
    "export_docx",
    # 审查清单服务
    "REVIEW_ITEMS_TEMPLATE",
    "complete_review",
    "create_default_review_items",
    "create_review",
    "get_review_by_report_id",
    "get_review_items_template",
    "get_review_statistics",
    "update_review_comments",
    "update_review_items",
]


# ---------------------------------------------------------------------------
# 人工审查清单服务
# ---------------------------------------------------------------------------


# 11项标准化审查项模板
REVIEW_ITEMS_TEMPLATE: dict[str, dict[str, Any]] = {
    "item_01": {
        "id": "item_01",
        "title": "事实认定审查",
        "description": "审查案件事实是否清楚，证据是否确实、充分",
        "default_checked": False,
        "category": "事实审查",
    },
    "item_02": {
        "id": "item_02",
        "title": "证据合法性审查",
        "description": "审查证据收集程序是否合法，是否存在非法证据排除情形",
        "default_checked": False,
        "category": "证据审查",
    },
    "item_03": {
        "id": "item_03",
        "title": "证据关联性审查",
        "description": "审查证据与案件事实之间是否存在关联性",
        "default_checked": False,
        "category": "证据审查",
    },
    "item_04": {
        "id": "item_04",
        "title": "主观明知认定审查",
        "description": "审查嫌疑人主观明知程度的认定是否合理",
        "default_checked": False,
        "category": "要件审查",
    },
    "item_05": {
        "id": "item_05",
        "title": "构成要件齐备性审查",
        "description": "审查帮信罪构成要件是否齐备",
        "default_checked": False,
        "category": "要件审查",
    },
    "item_06": {
        "id": "item_06",
        "title": "情节严重程度审查",
        "description": "审查情节档级（T1-T4）认定是否适当",
        "default_checked": False,
        "category": "量刑审查",
    },
    "item_07": {
        "id": "item_07",
        "title": "法律适用审查",
        "description": "审查法律条文引用是否准确，法律适用是否正确",
        "default_checked": False,
        "category": "法律审查",
    },
    "item_08": {
        "id": "item_08",
        "title": "量刑建议适当性审查",
        "description": "审查量刑建议是否在法定幅度内，是否适当",
        "default_checked": False,
        "category": "量刑审查",
    },
    "item_09": {
        "id": "item_09",
        "title": "程序合法性审查",
        "description": "审查办案程序是否符合法律规定",
        "default_checked": False,
        "category": "程序审查",
    },
    "item_10": {
        "id": "item_10",
        "title": "权利保障审查",
        "description": "审查嫌疑人诉讼权利是否得到保障",
        "default_checked": False,
        "category": "程序审查",
    },
    "item_11": {
        "id": "item_11",
        "title": "综合结论审查",
        "description": "审查综合分析结论是否合理，是否有遗漏或矛盾",
        "default_checked": False,
        "category": "综合审查",
    },
}


def get_review_items_template() -> dict[str, dict[str, Any]]:
    """获取审查项模板.

    Returns:
        dict: 11项审查项模板
    """
    return REVIEW_ITEMS_TEMPLATE.copy()


def create_default_review_items() -> dict[str, bool]:
    """创建默认审查项状态.

    Returns:
        dict: 所有审查项默认未勾选状态
    """
    return {
        item_id: item["default_checked"]
        for item_id, item in REVIEW_ITEMS_TEMPLATE.items()
    }


async def create_review(
    db: AsyncSession,
    report_id: int,
    reviewer_id: int | None = None,
) -> ReportReview:
    """创建审查记录.

    Args:
        db: 数据库会话
        report_id: 报告ID
        reviewer_id: 审查人ID

    Returns:
        ReportReview: 创建的审查记录
    """
    logger.info(f"创建审查记录 - 报告ID={report_id}, 审查人ID={reviewer_id}")

    review = ReportReview(
        report_id=report_id,
        reviewer_id=reviewer_id,
        items=create_default_review_items(),
        comments=None,
        completed_at=None,
    )

    db.add(review)
    await db.commit()
    await db.refresh(review)

    logger.info(f"审查记录创建成功 - 审查ID={review.id}")

    return review


async def get_review_by_report_id(
    db: AsyncSession,
    report_id: int,
) -> ReportReview | None:
    """根据报告ID获取审查记录.

    Args:
        db: 数据库会话
        report_id: 报告ID

    Returns:
        ReportReview | None: 审查记录或None
    """
    result = await db.execute(
        select(ReportReview).where(ReportReview.report_id == report_id)
    )
    return result.scalar_one_or_none()


async def update_review_items(
    db: AsyncSession,
    review_id: int,
    items: dict[str, bool],
) -> ReportReview | None:
    """更新审查项状态.

    Args:
        db: 数据库会话
        review_id: 审查记录ID
        items: 审查项状态字典

    Returns:
        ReportReview | None: 更新后的审查记录或None
    """
    logger.info(f"更新审查项状态 - 审查ID={review_id}")

    result = await db.execute(
        select(ReportReview).where(ReportReview.id == review_id)
    )
    review = result.scalar_one_or_none()

    if not review:
        logger.warning(f"审查记录不存在 - 审查ID={review_id}")
        return None

    # 更新审查项状态
    review.items = items
    await db.commit()
    await db.refresh(review)

    logger.info(f"审查项状态更新成功 - 审查ID={review_id}")

    return review


async def update_review_comments(
    db: AsyncSession,
    review_id: int,
    comments: str,
) -> ReportReview | None:
    """更新审查意见.

    Args:
        db: 数据库会话
        review_id: 审查记录ID
        comments: 审查意见

    Returns:
        ReportReview | None: 更新后的审查记录或None
    """
    logger.info(f"更新审查意见 - 审查ID={review_id}")

    result = await db.execute(
        select(ReportReview).where(ReportReview.id == review_id)
    )
    review = result.scalar_one_or_none()

    if not review:
        logger.warning(f"审查记录不存在 - 审查ID={review_id}")
        return None

    review.comments = comments
    await db.commit()
    await db.refresh(review)

    logger.info(f"审查意见更新成功 - 审查ID={review_id}")

    return review


async def complete_review(
    db: AsyncSession,
    review_id: int,
    items: dict[str, bool] | None = None,
    comments: str | None = None,
) -> ReportReview | None:
    """完成审查.

    更新审查项状态、审查意见，并设置审查完成时间。

    Args:
        db: 数据库会话
        review_id: 审查记录ID
        items: 审查项状态字典（可选）
        comments: 审查意见（可选）

    Returns:
        ReportReview | None: 更新后的审查记录或None
    """
    logger.info(f"完成审查 - 审查ID={review_id}")

    result = await db.execute(
        select(ReportReview).where(ReportReview.id == review_id)
    )
    review = result.scalar_one_or_none()

    if not review:
        logger.warning(f"审查记录不存在 - 审查ID={review_id}")
        return None

    # 更新审查项状态
    if items is not None:
        review.items = items

    # 更新审查意见
    if comments is not None:
        review.comments = comments

    # 设置审查完成时间
    review.completed_at = datetime.now()

    await db.commit()
    await db.refresh(review)

    logger.info(f"审查完成 - 审查ID={review_id}")

    return review


async def get_review_statistics(
    db: AsyncSession,
    report_id: int,
) -> dict[str, Any]:
    """获取审查统计信息.

    Args:
        db: 数据库会话
        report_id: 报告ID

    Returns:
        dict: 审查统计信息
    """
    review = await get_review_by_report_id(db, report_id)

    if not review:
        return {
            "total_items": len(REVIEW_ITEMS_TEMPLATE),
            "checked_items": 0,
            "unchecked_items": len(REVIEW_ITEMS_TEMPLATE),
            "completion_rate": 0.0,
            "is_completed": False,
        }

    items = review.items or {}
    checked_count = sum(1 for v in items.values() if v)
    total_count = len(REVIEW_ITEMS_TEMPLATE)

    return {
        "total_items": total_count,
        "checked_items": checked_count,
        "unchecked_items": total_count - checked_count,
        "completion_rate": checked_count / total_count if total_count > 0 else 0.0,
        "is_completed": review.completed_at is not None,
        "completed_at": review.completed_at.isoformat() if review.completed_at else None,
    }
